"""Helpers a sheet's content.py imports, plus the layout shims rw.py render uses.

Layout mechanics are lifted from the Kinematics reference build
(copy-revision-worksheet-with-different-practice/bank_worked_sheet.py) so every
sheet this skill produces lays out the way that one did: padded literal part
labels in Examples, hoisted first parts on parts-only Practice questions, and
literal labels for every part after a hoist (SQ's auto-numbering would restart).
"""
from __future__ import annotations
import re
import sys
from pathlib import Path

SKILLS = Path(__file__).resolve().parent.parent
for d in ("copy-revision-worksheet-with-different-practice", "create-worksheet"):
    p = SKILLS / d
    if str(p) not in sys.path:
        sys.path.insert(0, str(p))

import revision_lib as R                      # noqa: E402
from docx.shared import Cm, RGBColor           # noqa: E402
from docx.enum.text import WD_TAB_ALIGNMENT    # noqa: E402

GREY = RGBColor(0x7F, 0x7F, 0x7F)
LIGHT = RGBColor(0xA6, 0xA6, 0xA6)

T = lambda s: ('text', s)                                   # noqa: E731
B = lambda s: ('text', s, {'bold': True})                   # noqa: E731
I = lambda s: ('text', s, {'italic': True})                 # noqa: E731
M = lambda s: ('math', s)                                   # noqa: E731
P = lambda s: ('text', s, {'italic': True, 'color': GREY})  # noqa: E731  principle line


def sm(text) -> list:
    """Bank text -> parts. Bare-superscript fragments like ``ms$^{-1}$`` hand
    pandoc the invalid latex ``^{-1}``; give them an empty base first."""
    return R.split_math(str(text or "").replace("$^", "${}^"))


ROMANS = {"i": 1, "ii": 2, "iii": 3, "iv": 4, "v": 5, "vi": 6}


def sorted_parts(row) -> list:
    """Parts arrive scrambled sometimes — sort by label, roman-aware."""
    parts = [p for p in R._parts(row) if isinstance(p, dict)]
    labels = [str(p.get("label") or "").strip().lower() for p in parts]
    if parts and all(l in ROMANS for l in labels):
        return sorted(parts, key=lambda p: ROMANS[str(p["label"]).strip().lower()])
    return sorted(parts, key=lambda p: str(p.get("label") or "").strip().lower())


def pad_label(lab) -> str:
    lab = str(lab).strip().strip("()")
    width = 6 if lab.lower() in ROMANS else 4
    return f"({lab})".ljust(width)


def lit_part(ws, lab, parts, marks=None):
    """A Practice part with a LITERAL label at 1.0 cm, text at 2.0 cm — the
    geometry SQ produces natively, without SQ's per-question restart."""
    p = ws.para([("text", f"({str(lab).strip().strip('()')})\t")] + parts, marks=marks)
    pf = p.paragraph_format
    pf.left_indent = Cm(2.0)
    pf.first_line_indent = Cm(-1.0)
    pf.tab_stops.add_tab_stop(Cm(2.0), WD_TAB_ALIGNMENT.LEFT)
    return p


def hoist_Q(ws, lab, parts, marks=None):
    """Parts-only question (empty stem): the first part rides the number line
    as ``1.  (a)  text`` instead of leaving the auto-number alone."""
    p = ws.Q([("text", f"({str(lab).strip().strip('()')})\t")] + parts, marks=marks)
    pf = p.paragraph_format
    pf.left_indent = Cm(2.0)
    pf.first_line_indent = Cm(-2.0)
    pf.tab_stops.add_tab_stop(Cm(1.0), WD_TAB_ALIGNMENT.LEFT)
    pf.tab_stops.add_tab_stop(Cm(2.0), WD_TAB_ALIGNMENT.LEFT)
    return p


def normalise_figure(fig: dict, stem: Path):
    """Re-encode a stored figure to PNG. python-docx only parses JFIF/EXIF JPEG
    headers and several bank JPEGs start with the quantisation table; Pillow
    reads them all. None if even Pillow cannot — the question is then laid out
    without its diagram rather than aborting the sheet."""
    from io import BytesIO
    out = stem.with_suffix(".png")
    try:
        from PIL import Image
        img = Image.open(BytesIO(fig["bytes"]))
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        img.save(out, format="PNG")
        return out
    except Exception:
        return None


# A figure is bounded on BOTH sides, aspect ratio intact. Capping width alone let a
# portrait diagram (590 x 771 px) render 10.5 cm wide and 13.7 cm TALL — two-thirds of
# a page for one practice question (Adrian, 5 Sep 2026). The height cap is what makes a
# scan "a suitable size"; worksheet_lib never upscales past the image's natural size, so
# a small render still stays small.
FIG_MAX_W = 10.5      # cm — the text column is 16 cm; wider reads as a plate, not a figure
FIG_MAX_H_EXAMPLE = 6.5   # glued to a solution box, so it has to leave room for one
FIG_MAX_H_PRACTICE = 7.5  # stands alone under its question
# A tall portrait scan — a graph-paper grid is the usual one — hits the height cap first
# and comes out 4.3 cm wide, too small to read the axis labels on. Give every figure a
# floor on WIDTH, and let it grow past the soft height cap up to a hard ceiling to honour
# that floor. 9.5 cm is about a third of the text height: big enough to read, small
# enough that it never takes the page over.
FIG_MIN_W = 6.0
FIG_HARD_H = 9.5


def figure_width_cm(fig: dict, max_h: float = FIG_MAX_H_PRACTICE) -> float:
    """Width in cm that keeps the figure inside the caps at its own aspect ratio."""
    px = fig.get("px") or (0, 0)
    w_px, h_px = (px[0] or 600), (px[1] or 400)
    ratio = (h_px / w_px) if w_px else 1.0
    w = min(FIG_MAX_W, w_px / 96 * 2.54)
    if w * ratio > max_h:
        w = max_h / ratio
    if w < FIG_MIN_W:
        w = min(FIG_MIN_W, FIG_HARD_H / ratio)
    return round(w, 2)


# ---------------------------------------------------------------------------
# Data tables inside a question's text
# ---------------------------------------------------------------------------
# Extractor generations stored tables two ways and pandoc renders NEITHER as a
# table: a markdown pipe run collapses to literal "| | Mean mass | Standard
# deviation | | Papayas | 1.85 |" text (Adrian, 5 Sep 2026), and a LaTeX array
# becomes a borderless OMML matrix. Both are turned into real Word tables here.
# The markdown form has no newlines — rows run together, so the row boundary is
# the "| |" between a row's closing pipe and the next row's opening one.

MD_TABLE_RE = re.compile(r"(?:\|[^|\n]*){3,}\|")
TEX_TABLE_RE = re.compile(
    r"\$\$\s*\\begin\{(?P<e>array|tabular)\}(?:\{[^}]*\})?(?P<body>.*?)\\end\{(?P=e)\}\s*\$\$"
    r"|\\begin\{(?P<e2>array|tabular)\}(?:\{[^}]*\})?(?P<body2>.*?)\\end\{(?P=e2)\}",
    re.S)
_RULE_CELL = re.compile(r"^:?-{2,}:?$")


def _pad(rows):
    """A row that lost a leading EMPTY cell to the row-split is short by one."""
    rows = [r for r in rows if r and not all(_RULE_CELL.match(c) for c in r)]
    if len(rows) < 2:
        return None
    w = max(len(r) for r in rows)
    return [[""] * (w - len(r)) + r for r in rows] if w >= 2 else None


def md_table_rows(seg: str):
    rows = []
    for chunk in re.split(r"\|\s*\|", seg.strip()):
        cells = [c.strip() for c in chunk.split("|")]
        while cells and cells[0] == "":
            cells.pop(0)
        while cells and cells[-1] == "":
            cells.pop()
        if cells:
            rows.append(cells)
    return _pad(rows)


def tex_table_rows(body: str):
    body = body.replace(r"\hline", " ").replace(r"\\[", r"\\ [")
    rows = []
    for line in re.split(r"\\\\", body):
        if not line.strip():
            continue
        rows.append([c.strip() for c in line.split("&")])
    return _pad(rows)


def _table_segments(text: str):
    """[(kind, payload), …] with kind 'text' | 'md' | 'tex', in reading order."""
    spans = []
    for m in TEX_TABLE_RE.finditer(text):
        body = m.group("body") if m.group("body") is not None else m.group("body2")
        rows = tex_table_rows(body)
        if rows:
            spans.append((m.start(), m.end(), "tex", rows))
    for m in MD_TABLE_RE.finditer(text):
        if any(s < m.end() and m.start() < e for s, e, _, _ in spans):
            continue
        rows = md_table_rows(m.group(0))
        if rows:
            spans.append((m.start(), m.end(), "md", rows))
    spans.sort()
    out, pos = [], 0
    for s, e, kind, rows in spans:
        if text[pos:s].strip():
            out.append(("text", text[pos:s]))
        out.append((kind, rows))
        pos = e
    if text[pos:].strip():
        out.append(("text", text[pos:]))
    return out or [("text", text)]


def data_table(ws, rows, as_math=False):
    """A real Word table, bordered, centred cells, sized to the text column."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    ncol = max(len(r) for r in rows)
    t = ws.doc.add_table(rows=len(rows), cols=ncol)
    t.style = ws.doc.styles["Table Grid"]
    t.autofit = True
    for trow, row in zip(t.rows, rows):
        for cell, val in zip(trow.cells, list(row) + [""] * (ncol - len(row))):
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = p.paragraph_format.space_after = None
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if val.strip():
                ws._fill(p, [("math", val)] if as_math else sm(val))
    ws.doc.add_paragraph()
    return t


def emit_text(ws, text, emitter=None, marks=None, lead=None):
    """Lay out question text, rendering any embedded data table as a real table.

    `emitter(parts, marks=…)` places the FIRST prose chunk (so Q/SQ numbering and
    hanging indents still apply); later chunks are plain paragraphs. The mark
    allocation rides the last prose chunk — never a table, which cannot carry the
    right-aligned tab stop.
    """
    emitter = emitter or ws.para
    segs = _table_segments(str(text or ""))
    last_text = max((i for i, (k, _) in enumerate(segs) if k == "text"), default=-1)
    first_done = False
    for i, (kind, payload) in enumerate(segs):
        if kind == "text":
            parts = (lead or []) + sm(payload.strip()) if not first_done else sm(payload.strip())
            m = marks if i == last_text else None
            (emitter if not first_done else ws.para)(parts, marks=m)
            first_done = True
        else:
            data_table(ws, payload, as_math=(kind == "tex"))
    if not first_done:                      # text was nothing but a table
        emitter((lead or []) + [("text", "")], marks=marks)


# ---------------------------------------------------------------------------
# House header (Adrian's own revision sheets, read off them 5 Sep 2026)
# ---------------------------------------------------------------------------
# His sheets carry the title in the PAGE HEADER, not the body: two centred bold
# Times New Roman lines — the level line at 9.5 pt, the topic at 12 pt (the
# document default) — repeating on every page. There is no body title block and
# no footer. In Word's Print Layout the header text is dimmed grey; that is
# Word showing an inactive header, not a colour on the runs.

HDR_LEVEL_PT = 9.5
HDR_TOPIC_PT = 12.0


def _page_field(paragraph, instr: str):
    """A live Word field (PAGE / NUMPAGES) — python-docx has no helper for it."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for kind, text in (("begin", None), (None, instr), ("separate", None),
                       ("text", "1"), ("end", None)):
        r = OxmlElement("w:r")
        if kind == "text":
            t = OxmlElement("w:t"); t.text = text; r.append(t)
        elif kind is None:
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve"); it.text = text; r.append(it)
        else:
            fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), kind); r.append(fc)
        paragraph._p.append(r)


def house_header(ws, level_line: str, topic: str, page_numbers: bool = True):
    """Adrian's two-line running head, plus an optional page-number footer."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
    from docx.shared import Cm as _Cm

    sec = ws.doc.sections[0]
    sec.header_distance = _Cm(1.27)
    sec.footer_distance = _Cm(1.27)
    hdr = sec.header
    hdr.is_linked_to_previous = False
    paras = list(hdr.paragraphs)
    while len(paras) < 2:
        paras.append(hdr.add_paragraph())
    for extra in paras[2:]:
        extra._element.getparent().remove(extra._element)
    for para, text, size in ((paras[0], level_line, HDR_LEVEL_PT),
                             (paras[1], topic, HDR_TOPIC_PT)):
        for r in list(para.runs):
            r._element.getparent().remove(r._element)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = True

    if page_numbers:
        ftr = sec.footer
        ftr.is_linked_to_previous = False
        p = ftr.paragraphs[0]
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for txt, instr in (("Page ", None), (None, " PAGE "),
                           (" of ", None), (None, " NUMPAGES ")):
            if instr:
                _page_field(p, instr)
            else:
                run = p.add_run(txt)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
    return hdr
