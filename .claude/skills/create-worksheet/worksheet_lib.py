"""
worksheet_lib.py — single helper module for building Adrian's math worksheets.

Usage:
  1. Copy this file (or import from the skill folder) alongside your worksheet author script.
  2. In your author script, import the helpers and write your questions:

        from worksheet_lib import Worksheet

        ws = Worksheet()
        ws.title('My Worksheet')
        ws.subtitle('IP4 / Sec 4 Mathematics')

        ws.Q([('text', '(Topic)  Find '), ('math', 'A^2'), ('text', '.')], marks=3)
        ws.ans([('math', 'A^2 = ...')])

        ws.save('my_worksheet.docx')

  3. Run: python3 author_script.py

Everything (reference style setup, numbering definitions, OMML conversion,
inline numPr patching) is handled internally. Only one script to run.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import subprocess, tempfile, os, zipfile, io

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

NUMBERING_XML = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:abstractNum w:abstractNumId="100">
    <w:multiLevelType w:val="singleLevel"/>
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
      <w:pPr><w:ind w:left="567" w:hanging="567"/></w:pPr>
    </w:lvl>
  </w:abstractNum>
  <w:abstractNum w:abstractNumId="101">
    <w:multiLevelType w:val="singleLevel"/>
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerLetter"/>
      <w:lvlText w:val="(%1)"/>
      <w:lvlJc w:val="left"/>
      <w:pPr><w:ind w:left="1134" w:hanging="567"/></w:pPr>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1"><w:abstractNumId w:val="100"/></w:num>
'''
# 30 sub-question numIds, each with startOverride so (a)(b)(c) restarts per question
for _i in range(30):
    NUMBERING_XML += (
        f'  <w:num w:numId="{10+_i}">'
        f'<w:abstractNumId w:val="101"/>'
        f'<w:lvlOverride w:ilvl="0"><w:startOverride w:val="1"/></w:lvlOverride>'
        f'</w:num>\n'
    )
NUMBERING_XML += '</w:numbering>'


def _latex_to_omml(latex_expr, display=False):
    """Convert a LaTeX math expression to an OMML element via pandoc."""
    md = f"$${latex_expr}$$" if display else f"${latex_expr}$"
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
        f.write(md)
        md_path = f.name
    docx_path = md_path.replace('.md', '.docx')
    try:
        subprocess.run(['pandoc', md_path, '-o', docx_path], check=True, capture_output=True)
        with zipfile.ZipFile(docx_path) as z:
            doc_xml = z.read('word/document.xml')
        tree = etree.fromstring(doc_xml)
        if display:
            elem = tree.find(f'.//{{{M_NS}}}oMathPara')
            if elem is None:
                elem = tree.find(f'.//{{{M_NS}}}oMath')
        else:
            elem = tree.find(f'.//{{{M_NS}}}oMath')
        return elem
    finally:
        os.unlink(md_path)
        if os.path.exists(docx_path):
            os.unlink(docx_path)


def _outer_border_only(table):
    """Adrian's solution boxes show ONLY the outer TableGrid border.

    Instance-level override copied from his real Revision sheets: insideH and
    insideV set to none, outer edges inherited from the TableGrid style.
    """
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        borders.append(el)
    anchor = None  # CT_TblPr schema order: tblBorders sits before these three
    for tag in ('w:tblLayout', 'w:tblCellMar', 'w:tblLook'):
        anchor = tblPr.find(qn(tag))
        if anchor is not None:
            break
    if anchor is not None:
        anchor.addprevious(borders)
    else:
        tblPr.append(borders)


def _cant_split(row):
    """Forbid Word from splitting this table row across a page."""
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn('w:cantSplit')) is None:
        trPr.append(OxmlElement('w:cantSplit'))


def _left_align_math(elem):
    """Point an oMathPara's own justification left (pandoc emits center)."""
    if elem is None or elem.tag != f'{{{M_NS}}}oMathPara':
        return
    pr = elem.find(f'{{{M_NS}}}oMathParaPr')
    if pr is None:
        pr = etree.SubElement(elem, f'{{{M_NS}}}oMathParaPr')
        elem.insert(0, pr)  # oMathParaPr must be the first child
    jc = pr.find(f'{{{M_NS}}}jc')
    if jc is None:
        jc = etree.SubElement(pr, f'{{{M_NS}}}jc')
    jc.set(f'{{{M_NS}}}val', 'left')


def _style_annotations(elem):
    """Grey out Adrian's '←' step annotations inside converted OMML.

    His sheets end a working line with a small grey note ("← apply chain
    rule"). Authors write it as \\quad\\text{← ...} in the latex; here every
    math run from the arrow onwards gets his exact styling: 50%-grey
    (7F7F7F), 8 pt.
    """
    if elem is None:
        return
    for mt in elem.iter(f'{{{M_NS}}}t'):
        if '←' not in (mt.text or ''):
            continue
        arrow_run = mt.getparent()
        parent = arrow_run.getparent()
        seen = False
        for sib in list(parent):
            if sib is arrow_run:
                seen = True
            if not seen or sib.tag != f'{{{M_NS}}}r':
                continue
            wrpr = sib.find(qn('w:rPr'))
            if wrpr is None:
                wrpr = OxmlElement('w:rPr')
                mrpr = sib.find(f'{{{M_NS}}}rPr')
                if mrpr is not None:
                    mrpr.addnext(wrpr)
                else:
                    sib.insert(0, wrpr)
            for tag in ('w:rFonts', 'w:color', 'w:sz', 'w:szCs'):
                old = wrpr.find(qn(tag))
                if old is not None:
                    wrpr.remove(old)
            rf = OxmlElement('w:rFonts')
            rf.set(qn('w:ascii'), 'Cambria Math')
            rf.set(qn('w:hAnsi'), 'Cambria Math')
            col = OxmlElement('w:color')
            col.set(qn('w:val'), '7F7F7F')
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '16')
            szc = OxmlElement('w:szCs')
            szc.set(qn('w:val'), '16')
            for e in (rf, col, sz, szc):
                wrpr.append(e)


class Worksheet:
    """Builder for a single worksheet docx with Adrian's house style."""

    #: one body line at 9.5 pt on 1.5 spacing
    LINE_PT = 9.5 * 1.5

    def __init__(self, working_space=0.0, keep_questions_together=False,
                 keep_figures_with_text=True):
        """working_space: blank writing lines to leave per mark, after every
        paragraph that carries a mark allocation. 0 disables it (the right
        choice for a solutions sheet); pass 2.5 for a worksheet students write
        on, so a [3] gets three times the room of a [1].

        keep_questions_together: glue a whole question so it cannot straddle a
        page break. OFF by default, and leave it off once working_space is
        generous — a question is then taller than a page, so Word either splits
        it anyway or bumps it whole and wastes most of the previous page (which
        is what left a title-only first page). What stops the ugly breaks is
        unit-level gluing, always on: a question's text keeps with its writing
        space, the answer line keeps with the line above it, and a figure keeps
        with its stem. Pages then break between blank writing lines, where a
        break costs nothing.

        keep_figures_with_text: anchor each figure to the paragraphs either
        side of it, so a diagram is never split from the stem that introduces
        it."""
        self.doc = Document()
        self.working_space = float(working_space)
        self.keep_questions_together = bool(keep_questions_together)
        self.keep_figures_with_text = bool(keep_figures_with_text)
        self._auto_subq_id = 9   # increments to 10, 11, ... per Q with sub-parts
        self._current_subq_id = None
        self._block_paras = []   # paragraphs of the current question block (for keep-together)
        self._blocks = []        # every finished block, for block_heights()
        self._fig_cm = {}        # id(paragraph) -> rendered figure height in cm
        self._example_n = 0      # auto-counter for example() labels
        self._setup_page()
        self._setup_styles()

    # ---------- private setup ----------
    def _setup_page(self):
        s = self.doc.sections[0]
        s.page_width = Cm(21)
        s.page_height = Cm(29.7)
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(1)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    def _setup_styles(self):
        # Normal: TNR 9.5pt 1.5 line spacing
        n = self.doc.styles['Normal']
        n.font.name = 'Times New Roman'
        n.font.size = Pt(9.5)
        n.paragraph_format.line_spacing = 1.5
        n.paragraph_format.space_before = Pt(0)
        n.paragraph_format.space_after = Pt(0)
        rPr = n.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        for attr in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
            rFonts.set(qn(f'w:{attr}'), 'Times New Roman')

        # WSTitle
        t = self.doc.styles.add_style('WSTitle', WD_STYLE_TYPE.PARAGRAPH)
        t.base_style = self.doc.styles['Normal']
        t.font.name = 'Times New Roman'
        t.font.size = Pt(12)
        t.font.bold = True
        t.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        t.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t.paragraph_format.space_after = Pt(6)

        # WSSubtitle
        st = self.doc.styles.add_style('WSSubtitle', WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = self.doc.styles['Normal']
        st.font.name = 'Times New Roman'
        st.font.size = Pt(10)
        st.font.italic = True
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        st.paragraph_format.space_after = Pt(8)

        # Answer
        a = self.doc.styles.add_style('Answer', WD_STYLE_TYPE.PARAGRAPH)
        a.base_style = self.doc.styles['Normal']
        a.font.color.rgb = RGBColor(0x84, 0x3C, 0x0C)
        a.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Question, SubQuestion (no inline numbering — applied per paragraph)
        self.doc.styles.add_style('Question', WD_STYLE_TYPE.PARAGRAPH).base_style = self.doc.styles['Normal']
        self.doc.styles.add_style('SubQuestion', WD_STYLE_TYPE.PARAGRAPH).base_style = self.doc.styles['Normal']

        # Set zoom
        zoom = self.doc.settings.element.find(qn('w:zoom'))
        if zoom is None:
            z = etree.SubElement(self.doc.settings.element, qn('w:zoom'))
            z.set(qn('w:percent'), '100')
        else:
            zoom.set(qn('w:percent'), '100')

    # ---------- private paragraph builder ----------
    def _add(self, parts, style=None, num_id=None, marks=None, alignment=None):
        p = self.doc.add_paragraph()
        if style:
            p.style = self.doc.styles[style]
        if alignment is not None:
            p.alignment = alignment

        if num_id is not None:
            pPr = p._element.get_or_add_pPr()
            existing = pPr.find(qn('w:numPr'))
            if existing is not None:
                pPr.remove(existing)
            numPr = OxmlElement('w:numPr')
            ilvl = OxmlElement('w:ilvl')
            ilvl.set(qn('w:val'), '0')
            nId = OxmlElement('w:numId')
            nId.set(qn('w:val'), str(num_id))
            numPr.append(ilvl)
            numPr.append(nId)
            pPr.append(numPr)

        self._fill(p, parts)

        if marks is not None:
            p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT)
            run = p.add_run(f'\t[{marks}]')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9.5)
        self._block_paras.append(p)
        if marks is not None and self.working_space:
            self.workspace(marks=marks)
        return p

    def _fill(self, p, parts):
        """Append runs/math to an existing paragraph from a parts list."""
        for part in parts:
            kind = part[0]
            if kind == 'text':
                text = part[1]
                attrs = part[2] if len(part) > 2 else {}
                run = p.add_run(text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(9.5)
                if attrs.get('bold'):
                    run.bold = True
                if attrs.get('italic'):
                    run.italic = True
                if attrs.get('color'):
                    run.font.color.rgb = attrs['color']
            elif kind == 'math':
                elem = _latex_to_omml(part[1], display=False)
                if elem is not None:
                    _style_annotations(elem)
                    p._element.append(elem)
            elif kind == 'math_display':
                elem = _latex_to_omml(part[1], display=True)
                if elem is not None:
                    _style_annotations(elem)
                    p._element.append(elem)

    # ---------- public API ----------
    def title(self, text):
        p = self.doc.add_paragraph()
        p.style = self.doc.styles['WSTitle']
        p.add_run(text)

    def subtitle(self, text):
        p = self.doc.add_paragraph()
        p.style = self.doc.styles['WSSubtitle']
        p.add_run(text)

    def concept(self, text):
        """Bold concept subtitle written above the Example(s) it covers.

        Adrian's convention: the concept a worked example teaches is a bold
        line of its own, in front — then the "Example N" label, then the
        question. When several related examples share one concept, write the
        concept once and give the examples one number with letters
        (example('a'), example('b'))."""
        return self._add([('text', text, {'bold': True})])

    def example(self, letter=None):
        """Bold auto-numbered "Example N" label line.

        example()    -> Example 1, Example 2, ... (counter increments)
        example('a') -> starts a lettered group: Example 3a (counter increments)
        example('b') -> Example 3b (same number as the last 'a')
        """
        if letter in (None, '', 'a'):
            self._example_n += 1
        return self._add([('text', f'Example {self._example_n}{letter or ""}',
                           {'bold': True})])

    def Q(self, parts, marks=None):
        """Main question. Auto-numbered 1. 2. 3. ..."""
        # Bump the sub-question id pool for this question; reset on each Q call
        self._finish_block()    # glue the question that just ended
        self._auto_subq_id += 1
        self._current_subq_id = self._auto_subq_id
        self._block_paras = []  # a new question starts a new keep-together block
        # Apply inline numId=1 directly so MS Word picks it up reliably
        return self._add(parts, style='Question', num_id=1, marks=marks)

    def SQ(self, parts, marks=None):
        """Sub-question (a)(b)(c) ... auto-tracks under the current main question."""
        if self._current_subq_id is None:
            raise RuntimeError('SQ() called before any Q(). Add a main question first.')
        return self._add(parts, style='SubQuestion',
                         num_id=self._current_subq_id, marks=marks)

    def para(self, parts, marks=None):
        """Plain paragraph (no numbering)."""
        return self._add(parts, marks=marks)

    def section(self, text):
        """Bold section header, e.g. 'Section B - Congruency and Similarity'.

        Closes the previous question and glues itself to the question that
        follows, so a header can never be left stranded at the foot of a page
        with its first question overleaf."""
        self._finish_block()
        p = self._add([('text', text, {'bold': True})])
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.keep_with_next = True
        return p

    def math_block(self, latex_expr):
        """Centred display equation."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elem = _latex_to_omml(latex_expr, display=True)
        if elem is not None:
            _style_annotations(elem)
            p._element.append(elem)
        self._block_paras.append(p)
        return p

    def ans(self, parts):
        """[Ans: ...] right-aligned orange line, glued to the line above it so
        it cannot be stranded at the top of the next page."""
        if self._block_paras:
            self._block_paras[-1].paragraph_format.keep_with_next = True
        full = [('text', '[Ans: ')] + list(parts) + [('text', ']')]
        return self._add(full, style='Answer', alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    def _picture(self, p, path, width_cm):
        """Centre a PNG in paragraph p, capped at width_cm and never upscaled
        past the image's natural 96-dpi size — a small render should stay
        small, not blur."""
        from PIL import Image  # python-docx already depends on Pillow
        with Image.open(path) as im:
            natural_cm = im.width / 96 * 2.54
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        w = min(width_cm, 16, natural_cm)
        run.add_picture(path, width=Cm(w))
        with Image.open(path) as im:
            self._fig_cm[id(p)] = w * im.height / im.width
        return p

    def figure(self, path, width_cm=10.5):
        """Embed a rendered figure (see figure_lib.py) under the current question."""
        p = self._picture(self.doc.add_paragraph(), path, width_cm)
        if self.keep_figures_with_text:
            if self._block_paras:          # stay with the stem above
                self._block_paras[-1].paragraph_format.keep_with_next = True
            p.paragraph_format.keep_with_next = True   # and the part below
        self._block_paras.append(p)
        return p

    def solution_box(self, rows, keep_together=True):
        """Boxed worked solution in Adrian's house format.

        His Revision "(With Worked Examples)" sheets put every solution in a
        TableGrid table showing ONLY the outer border (no inner gridlines),
        under a bold "Solution:" line with one blank line of breathing space
        after the question. Two columns: the part label alone in a narrow
        first column, the working beside it, one table row per part. A blank
        line is inserted automatically between parts (never after the last).

        rows: list of (label, steps). label is '(a)' / '(i)' ('' for an
        unlabelled single-cell solution). Each step is one of:

        - a bare latex string — a display equation, left-aligned at a small
          indent. Never chain several = signs on one line: write multi-step
          working as a \\begin{aligned} block (&= per line) so the lines
          stack with the = signs vertically aligned. End a line with
          \\quad\\text{← short note} for Adrian's grey 8pt arrow annotation.
        - a parts list (same shapes Q()/para() take), rendered left-aligned
          — for prose steps and inline-math sentences.
        - ('figure', path[, width_cm]) — a centred figure_lib PNG inside the
          box, for a sketch that explains the working (default 8 cm).

        keep_together (default True) keeps the question paragraphs (since
        the last Q()), the "Solution:" line and the whole box on one page —
        Word pushes the block to a fresh page rather than straddling. A block
        taller than a full page still splits gracefully.
        """
        spacer = self.doc.add_paragraph()  # breathing space above "Solution:"
        self._block_paras.append(spacer)
        self._add([('text', 'Solution:', {'bold': True})])
        labelled = any(label for label, _ in rows)
        table = self.doc.add_table(rows=len(rows), cols=2 if labelled else 1)
        table.style = self.doc.styles['Table Grid']
        table.autofit = False
        _outer_border_only(table)
        for idx, ((label, steps), row) in enumerate(zip(rows, table.rows)):
            if labelled:
                lab_cell, work_cell = row.cells
                lab_cell.width = Cm(1.0)
                work_cell.width = Cm(15.0)
                if label:
                    self._fill(lab_cell.paragraphs[0], [('text', label)])
            else:
                work_cell = row.cells[0]
                work_cell.width = Cm(16.0)
            first = True
            for step in steps:
                p = work_cell.paragraphs[0] if first else work_cell.add_paragraph()
                first = False
                p.paragraph_format.line_spacing = 1.15  # boxes are tighter than the 1.5 body
                if isinstance(step, tuple) and step and step[0] == 'figure':
                    self._picture(p, step[1], step[2] if len(step) > 2 else 8.0)
                elif isinstance(step, str):
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.left_indent = Cm(0.5)
                    elem = _latex_to_omml(step, display=True)
                    if elem is not None:
                        _left_align_math(elem)
                        _style_annotations(elem)
                        p._element.append(elem)
                else:
                    self._fill(p, step)
            if idx < len(rows) - 1:  # one blank line between parts, none after the last
                gap = work_cell.add_paragraph()
                gap.paragraph_format.line_spacing = 1.15
        if keep_together:
            for para in self._block_paras:
                para.paragraph_format.keep_with_next = True
            for row in table.rows:
                _cant_split(row)
            for row in list(table.rows)[:-1]:  # last row must NOT keep with what follows
                for cell in row.cells:
                    for cp in cell.paragraphs:
                        cp.paragraph_format.keep_with_next = True
        self._block_paras = []
        self.doc.add_paragraph()  # breathing space between the box and what follows
        return table

    def workspace(self, marks=None, lines=None):
        """Blank writing space: `marks` x self.working_space real empty lines
        (or an explicit `lines`). Emitted automatically after every marked
        paragraph when the Worksheet was built with working_space.

        Real empty paragraphs, NOT one paragraph with a big space_after: Word
        discards trailing space at a page break, so a space_after gap that
        straddles a break silently loses the rest of the student's writing room.
        Separate lines simply flow onto the next page.

        The marked paragraph and ALL of its blank lines are glued into one unit,
        so a part's writing space never straddles a page break — a page breaks
        between parts, never inside one. Keep the unit small: it is one line of
        text plus `marks x working_space` lines, so at 4.0 a [3] part is 13
        lines (~6.5 cm) and packs easily.
        """
        if lines is None:
            if marks is None:
                raise ValueError('workspace() needs marks or lines')
            lines = float(marks) * self.working_space
        n = int(round(lines))
        if n <= 0:
            return None
        if self._block_paras:                      # text keeps with its space
            self._block_paras[-1].paragraph_format.keep_with_next = True
        made = []
        for i in range(n):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            if i < n - 1:      # the whole run is atomic; only its end may break
                p.paragraph_format.keep_with_next = True
            self._block_paras.append(p)
            made.append(p)
        return made

    def _finish_block(self):
        """Close the question just finished. The last paragraph must NOT keep
        with what follows, or every question chains into one unbreakable block."""
        if self.keep_questions_together:
            for para in self._block_paras[:-1]:
                para.paragraph_format.keep_with_next = True
        if self._block_paras:
            self._blocks.append(self._block_paras)
        self._block_paras = []

    #: usable text column, A4 less the house margins (top 2 cm, bottom 1 cm)
    PAGE_CM = 26.7
    LINE_CM = LINE_PT / 28.35
    CHAR_CM = 0.151        # Times 9.5 pt, averaged over mixed case

    def block_heights(self):
        """Estimated height, in cm, of each question block — the check that
        keep_questions_together is actually achievable. A block taller than
        PAGE_CM cannot be kept whole; one just under it will be bumped to a
        fresh page and leave the previous one mostly blank.

        An estimate, not a measurement: line counts come from character counts
        at the body size, and an equation is charged as three characters. Treat
        anything within ~1.5 cm of PAGE_CM as "may not fit"."""
        return [round(self._height_cm(b), 1) for b in
                self._blocks + ([self._block_paras] if self._block_paras else [])]

    def _height_cm(self, paras):
        """Estimated height in cm of a list of paragraphs."""
        if True:
            cm = 0.0
            for p in paras:
                if id(p) in self._fig_cm:
                    cm += self._fig_cm[id(p)] + 8 / 28.35   # picture + its 4pt padding
                    continue
                indent = 1.5 if p.style is not None and p.style.name == 'SubQuestion' else 0.0
                width = 16.0 - indent
                chars = len(p.text) + 3 * p._element.xml.count('<m:oMath')
                cm += max(1, -(-chars // int(width / self.CHAR_CM))) * self.LINE_CM
            return cm

    def glued_heights(self):
        """Height in cm of each atomic run — a chain of keep_with_next
        paragraphs plus the one that ends it. These, not whole questions, are
        what must fit on a page: Word has to break inside any run taller than
        PAGE_CM. Same estimating caveats as block_heights()."""
        flat = [p for block in self._blocks + ([self._block_paras] if self._block_paras else [])
                for p in block]
        runs, cur = [], []
        for p in flat:
            cur.append(p)
            if not p.paragraph_format.keep_with_next:
                runs.append(cur)
                cur = []
        if cur:
            runs.append(cur)
        return [round(self._height_cm(run), 1) for run in runs]

    def page_break(self):
        self._finish_block()
        self.doc.add_page_break()
        self._block_paras = []  # a manual break ends any keep-together block

    def save(self, path):
        """Save the worksheet, injecting clean numbering.xml."""
        self._finish_block()    # the last question has no Q() after it
        # Save to a temp buffer first, then rewrite numbering.xml
        buf = io.BytesIO()
        self.doc.save(buf)
        buf.seek(0)

        with zipfile.ZipFile(buf, 'r') as zin:
            items = {name: zin.read(name) for name in zin.namelist()}

        items['word/numbering.xml'] = NUMBERING_XML.encode('utf-8')

        # Schema fix: pandoc emits <m:mcJc> before <m:count> inside <m:mcPr>,
        # but OOXML requires count first (matrices fail strict validation otherwise).
        import re as _re
        items['word/document.xml'] = _re.sub(
            rb'(<m:mcJc m:val="[^"]*"/>)(<m:count m:val="[^"]*"/>)',
            rb'\2\1',
            items['word/document.xml'])

        with zipfile.ZipFile(path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for name, data in items.items():
                zout.writestr(name, data)

        print(f'Saved {path}')
