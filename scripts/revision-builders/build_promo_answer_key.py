# -*- coding: utf-8 -*-
"""JC1 H2 Mathematics promo answer key — answers only, no provenance.

Deliberately carries no school name and no year, as asked.
Hanging-indent numbered list: bold number at the margin, answer text at
1.2 cm, bold part labels, real Word equations via worksheet_lib's OMML path.
"""
from worksheet_lib import Worksheet
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn

NAVY = RGBColor(0x1F, 0x4E, 0x79)

ws = Worksheet()
ws.title('H2 Mathematics')
ws.subtitle('Promotional Examination  —  Answer Key')


def rule():
    p = ws.doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    from docx.oxml import OxmlElement
    bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '808080')
    bdr.append(bottom)
    pPr.append(bdr)
    p.paragraph_format.space_after = Pt(8)


def K(n, parts):
    """One numbered answer entry, hanging-indented under its number."""
    p = ws.doc.add_paragraph()
    p.style = ws.doc.styles['Normal']
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.first_line_indent = Cm(-1.2)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f'{n}.'.ljust(4))
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9.5)
    r.bold = True
    for wt in r._r.findall(qn('w:t')):
        wt.set(qn('xml:space'), 'preserve')
    ws._fill(p, parts)
    return p


def B(label):
    return ('text', f'({label}) ', {'bold': True})


def t(s):
    return ('text', s)


def m(s):
    return ('math', s)


GAP = ('text', ' ')   # em space between parts

rule()

K(1, [B('a'), m(r'a=-1'), t(', '), m(r'b=3'), t(', '), m(r'c=6'), GAP,
      B('b'), m(r'\approx 1\,020\,000'), t(' (3 s.f.)')])

K(2, [m(r'\dfrac{1}{\sqrt{1-x^{2}}}')])

K(3, [B('a'), m(r'-\dfrac{2}{5}<x<2'), GAP,
      B('b'), m(r'x<\ln 2'), GAP,
      B('c'), m(r'x<-\dfrac{4}{5}'), t(' or '), m(r'x>4')])

K(4, [B('a'), m(r'x_{n}=16x_{n-1},\ n\geqslant 1'), GAP,
      B('b'), m(r'x_{n}=x_{0}(16)^{n}'), GAP,
      B('c'), m(r'\dfrac{n}{2}(n+1)'), GAP,
      B('d'), m(r'u_{n}=\dfrac{n}{2}(n+1)+1')])

K(5, [B('a'), m(r'y=4a^{2}-x^{2}'), GAP,
      B('b'), t('Shown'), GAP,
      B('c'), m(r'\dfrac{16\sqrt{3}}{9}a^{3}')])

K(6, [B('a'), m(r'\vec{AC}\cdot\vec{BC}=0'), t(', so '), m(r'AC'), t(' and '),
      m(r'BC'), t(' are perpendicular  (shown)'), GAP,
      B('b'), t('The length of projection of '), m(r'AC'), t(' onto '), m(r'OA'), GAP,
      B('c'), m(r'C'), t(', '), m(r'M'), t(' and '), m(r'R'), t(' are collinear  (shown)')])

K(7, [B('a'), m(r'P=\left(\dfrac{1}{\ln 2},\ 2^{1/\ln 2}\right)'), GAP,
      ('text', '(b)(i) ', {'bold': True}), t('Shown'), GAP,
      ('text', '(b)(ii) ', {'bold': True}),
      m(r'\dfrac{d\theta}{dt}=-\dfrac{400}{(4t)^{2}+(100-t)^{2}}')])

K(8, [t('Graph sketches — no numerical answer.')])

K(9, [('text', '(a)(i) ', {'bold': True}), t('Asymptotes '), m(r'x=5'), t(', '),
      m(r'y=x+9'), t('; turning points '), m(r'(-1.32,\,1.35)'), t(' max, '),
      m(r'(11.3,\,26.6)'), t(' min'), GAP,
      ('text', '(a)(ii) ', {'bold': True}), m(r'-5<x\leqslant -3.52'),
      t(' or '), m(r'x>5'), GAP,
      B('b'), m(r'y=\dfrac{4x^{2}+8x-5}{2x-5}-2')])

K(10, [B('a'), m(r'f^{-1}(x)=\dfrac{2x-2\lambda}{x+2}'), GAP,
       B('b'), m(r'0<x<2'), GAP,
       ('text', '(c)(i) ', {'bold': True}), t('Since '),
       m(r'R_{g}=\left[-\lambda,\,0\right)'), t(' is not contained in '),
       m(r'D_{f}=\left(-\lambda,\,2\right)'), t(', '), m(r'fg'),
       t(' does not exist'), GAP,
       ('text', '(c)(ii) ', {'bold': True}),
       m(r'fg(x)=-\dfrac{2x^{2}}{x^{2}-\lambda-2}'), t(', range '),
       m(r'(0,\,\lambda)')])

K(11, [B('a'), m(r'3y+4z=20'), GAP,
       B('b'), m(r'\dfrac{9}{5}'), GAP,
       B('c'), m(r'p=2'), t(', '), m(r'q=5'), GAP,
       B('d'), m(r'\mathbf{r}=(2,\,0,\,5)+\lambda(-2,\,4,\,-3)'),
       t(', '), m(r'\lambda\in\mathbb{R}'), GAP,
       B('e'), m(r'3.7^{\circ}'), t(' (0.0654 rad)')])

K(12, [('text', '(a)(i) ', {'bold': True}), t('Shown — total '), m(r'1560'),
       t(' km < 1600 km'), GAP,
       ('text', '(a)(ii) ', {'bold': True}), t('124 km'), GAP,
       ('text', '(b)(i) ', {'bold': True}), t('$5748.80'), GAP,
       ('text', '(b)(ii) ', {'bold': True}), t('Shown'), GAP,
       ('text', '(b)(iii) ', {'bold': True}), t('2031')])

ws.doc.add_paragraph()
note = ws.doc.add_paragraph()
nr = note.add_run('Total: 100 marks  (4, 4, 7, 8, 8, 8, 9, 9, 9, 10, 12, 12)')
nr.font.name = 'Times New Roman'
nr.font.size = Pt(8.5)
nr.italic = True
nr.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

ws.save('promo_answer_key.docx')
print('saved')
