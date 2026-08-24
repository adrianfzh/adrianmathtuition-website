# -*- coding: utf-8 -*-
"""Hire Purchase — Worked Examples (student reference).

Every scenario a hire-purchase question can ask for, one worked example each:
the instalment, the hire-purchase price, the interest rate, the principal /
cash price, the length of the agreement, the deposit percentage, and a
two-plan comparison.  House style: Times New Roman question text, and the
navy/Consolas solution box from Adrian's own revision sheets.
"""
from worksheet_lib import Worksheet
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x4E, 0x79)
BAR = 'DEEBF7'
EDGE = '2E5C8A'


def _shade(cell, fill):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:color'), 'auto')
    el.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(el)


def _box(table, color):
    b = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '6')
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)
        b.append(e)
    tblPr = table._tbl.tblPr
    anchor = None
    for tag in ('w:tblLayout', 'w:tblCellMar', 'w:tblLook'):
        anchor = tblPr.find(qn(tag))
        if anchor is not None:
            break
    (anchor.addprevious(b) if anchor is not None else tblPr.append(b))


def _fix_grid(t):
    """python-docx writes an equal-width tblGrid whatever the cell widths say."""
    widths = []
    for cell in t.rows[0].cells:
        tcPr = cell._tc.find(qn('w:tcPr'))
        tcW = tcPr.find(qn('w:tcW')) if tcPr is not None else None
        widths.append(tcW.get(qn('w:w')) if tcW is not None else None)
    grid = t._tbl.find(qn('w:tblGrid'))
    if grid is None or any(w is None for w in widths):
        return
    cols = grid.findall(qn('w:gridCol'))
    if len(cols) == len(widths):
        for c, w in zip(cols, widths):
            c.set(qn('w:w'), w)


def mono_box(ws, title, lines, size=8.5):
    """Navy-edged box with a pale blue title bar and a Consolas body."""
    t = ws.doc.add_table(rows=2, cols=1)
    t.autofit = False
    _box(t, EDGE)
    head = t.rows[0].cells[0]
    head.width = Cm(16)
    _shade(head, BAR)
    hp = head.paragraphs[0]
    hp.paragraph_format.line_spacing = 1.0
    hp.paragraph_format.space_before = Pt(2)
    hp.paragraph_format.space_after = Pt(2)
    hr = hp.add_run(title)
    hr.font.name = 'Calibri'
    hr.font.size = Pt(10.5)
    hr.bold = True
    hr.font.color.rgb = NAVY

    body = t.rows[1].cells[0]
    body.width = Cm(16)
    first = True
    for line in lines:
        bp = body.paragraphs[0] if first else body.add_paragraph()
        first = False
        bp.paragraph_format.line_spacing = 1.0
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after = Pt(0)
        r = bp.add_run(line if line else ' ')
        r.font.name = 'Consolas'
        r.font.size = Pt(size)
        if line.startswith('(') or line.endswith(':') and not line.startswith(' '):
            r.bold = True
        for wt in r._r.findall(qn('w:t')):
            wt.set(qn('xml:space'), 'preserve')
    _fix_grid(t)
    ws._block_paras = []
    ws.doc.add_paragraph()
    return t


def T(s):
    return [('text', s)]


def head(ws, text):
    p = ws.doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10.5)
    r.bold = True
    r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.keep_with_next = True
    ws._block_paras.append(p)
    return p


def qtext(ws, text, indent=1.0):
    p = ws.doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9.5)
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.keep_with_next = True
    ws._block_paras.append(p)
    return p


ws = Worksheet()
ws.title('Hire Purchase')
ws.subtitle('Worked Examples — Every Type of Calculation')

# ------------------------------------------------------------------ notes
mono_box(ws, 'What the words mean, and the five relationships', [
 "Cash price      what you pay if you settle the whole amount at once",
 "Deposit         paid up front, often a percentage of the cash price",
 "Balance, P      = cash price − deposit",
 "                THIS is what interest is charged on, NOT the cash price",
 "",
 "1.  Interest         I = P × R × T ÷ 100",
 "                     R = % per year,  T = number of YEARS",
 "",
 "2.  Total to repay   = P + I",
 "",
 "3.  Monthly          = (P + I) ÷ (12 × T)",
 "    instalment",
 "",
 "4.  Hire purchase    = deposit + (monthly instalment × 12 × T)",
 "    price",
 "",
 "5.  Extra paid       = hire purchase price − cash price   ( = I )",
 "",
 "Rearranging relationship 1 gives the three 'work backwards' forms:",
 "     R = 100I ÷ (P × T)      P = 100I ÷ (R × T)      T = 100I ÷ (P × R)",
 "",
 "Simple interest is charged on the ORIGINAL balance for the whole term —",
 "it does not fall as you pay the loan off.",
])

# ------------------------------------------------------- 1. the instalment
head(ws, 'Example 1  —  Finding the monthly instalment')
qtext(ws, 'A washing machine has a cash price of $4500. Mr Tan pays a deposit of 20% of the cash '
          'price and repays the balance over 3 years, at 6% per annum simple interest, in equal '
          'monthly instalments. Calculate his monthly instalment.')
mono_box(ws, 'Solution  —  Example 1', [
 "Deposit  = 20% × 4500 = $900",
 "Balance  P = 4500 − 900 = $3600      ← interest is on 3600, not 4500",
 "",
 "Interest I = 3600 × 6 × 3 ÷ 100",
 "           = $648",
 "Total to repay = 3600 + 648 = $4248",
 "",
 "Number of months = 3 × 12 = 36",
 "Monthly instalment = 4248 ÷ 36",
 "                   = $118",
])

# --------------------------------------------- 2. HP price and extra paid
head(ws, 'Example 2  —  Finding the hire purchase price and the extra paid')
qtext(ws, 'For the washing machine in Example 1, find the total hire purchase price, and how much '
          'more Mr Tan pays than the cash price.')
mono_box(ws, 'Solution  —  Example 2', [
 "Hire purchase price = deposit + (instalment × months)",
 "                    = 900 + 36(118)",
 "                    = 900 + 4248",
 "                    = $5148",
 "",
 "Extra paid = 5148 − 4500 = $648",
 "",
 "Notice the extra paid is exactly the interest. That is always true,",
 "because deposit + balance = cash price, so the only thing added on",
 "top of the cash price is the interest.",
])

# ------------------------------------------------------ 3. interest rate
head(ws, 'Example 3  —  Finding the interest rate')
qtext(ws, 'A laptop has a cash price of $2400. Ryan pays a deposit of $600 and repays the balance '
          'in 24 equal monthly instalments of $87. Calculate the rate of simple interest per annum.')
mono_box(ws, 'Solution  —  Example 3', [
 "Balance  P = 2400 − 600 = $1800",
 "Total repaid on instalments = 24 × 87 = $2088",
 "",
 "Interest I = 2088 − 1800 = $288     ← what was paid, less what was owed",
 "",
 "T = 24 months = 2 years             ← the formula needs YEARS",
 "",
 "R = 100I ÷ (P × T)",
 "  = 100(288) ÷ (1800 × 2)",
 "  = 28 800 ÷ 3600",
 "  = 8% per annum",
])

# ------------------------------------------- 4. principal and cash price
head(ws, 'Example 4  —  Finding the balance and the cash price')
qtext(ws, 'A motorcycle is bought on hire purchase. A deposit of $3000 is paid, and the balance is '
          'repaid over 4 years at 5% per annum simple interest in monthly instalments of $290. '
          'Calculate the cash price of the motorcycle.')
mono_box(ws, 'Solution  —  Example 4', [
 "Total repaid on instalments = 48 × 290 = $13 920",
 "",
 "Let the balance be $P.",
 "Total repaid = P + interest",
 "             = P + P × 5 × 4 ÷ 100",
 "             = P + 0.2P",
 "             = 1.2P                 ← factorise, do not guess P",
 "",
 "1.2P = 13 920",
 "   P = 13 920 ÷ 1.2",
 "     = $11 600",
 "",
 "Cash price = balance + deposit",
 "           = 11 600 + 3000",
 "           = $14 600",
])

# ------------------------------------------------------------- 5. time
head(ws, 'Example 5  —  Finding the length of the agreement')
qtext(ws, 'A sofa has a cash price of $3200. Mrs Lee pays a deposit of 25% and repays the balance '
          'at 7% per annum simple interest in monthly instalments of $114. Find the number of '
          'years of the agreement.')
mono_box(ws, 'Solution  —  Example 5', [
 "Deposit = 25% × 3200 = $800",
 "Balance P = 3200 − 800 = $2400",
 "",
 "Let the agreement run for T years.",
 "",
 "Total repaid = 12T × 114 = 1368T           ← from the instalments",
 "Total repaid = 2400 + 2400 × 7 × T ÷ 100",
 "             = 2400 + 168T                 ← from balance + interest",
 "",
 "1368T = 2400 + 168T",
 "1200T = 2400",
 "    T = 2 years",
])

# ------------------------------------------------------ 6. deposit as %
head(ws, 'Example 6  —  Finding the deposit as a percentage')
qtext(ws, 'A television has a cash price of $2500. A customer pays a deposit, then repays the '
          'balance at 8% per annum simple interest in 36 monthly instalments of $62. Express the '
          'deposit as a percentage of the cash price.')
mono_box(ws, 'Solution  —  Example 6', [
 "Total repaid on instalments = 36 × 62 = $2232",
 "T = 36 months = 3 years",
 "",
 "Let the balance be $P.",
 "P + P × 8 × 3 ÷ 100 = 2232",
 "            P(1 + 0.24) = 2232",
 "                  1.24P = 2232",
 "                      P = $1800",
 "",
 "Deposit = 2500 − 1800 = $700",
 "",
 "Deposit as a percentage = 700 ÷ 2500 × 100%",
 "                        = 28%",
])

# ------------------------------------------------------- 7. comparison
head(ws, 'Example 7  —  Comparing two hire purchase plans')
qtext(ws, 'A refrigerator has a cash price of $1800. Two hire purchase plans are offered.')
qtext(ws, 'Plan A:  a deposit of 10%, and the balance over 2 years at 9% per annum.', indent=1.6)
qtext(ws, 'Plan B:  a deposit of 25%, and the balance over 3 years at 6% per annum.', indent=1.6)
qtext(ws, 'Determine which plan costs less in total, and state one reason a buyer might still '
          'choose the other plan.')
mono_box(ws, 'Solution  —  Example 7', [
 "Plan A",
 "  Deposit = 10% × 1800 = $180",
 "  Balance = 1800 − 180 = $1620",
 "  I = 1620 × 9 × 2 ÷ 100 = $291.60",
 "  Total to repay = 1620 + 291.60 = $1911.60",
 "  Instalment = 1911.60 ÷ 24 = $79.65",
 "  Hire purchase price = 180 + 1911.60 = $2091.60",
 "",
 "Plan B",
 "  Deposit = 25% × 1800 = $450",
 "  Balance = 1800 − 450 = $1350",
 "  I = 1350 × 6 × 3 ÷ 100 = $243",
 "  Total to repay = 1350 + 243 = $1593",
 "  Instalment = 1593 ÷ 36 = $44.25",
 "  Hire purchase price = 450 + 1593 = $2043",
 "",
 "Plan B costs 2091.60 − 2043 = $48.60 less in total.",
 "",
 "But a buyer short of cash now might still take Plan A: its deposit is",
 "$180 rather than $450, and it is over in 2 years instead of 3.",
])

ws.page_break()

# ------------------------------------------------------------- practice
head(ws, 'Practice')
qtext(ws, 'Answers are at the foot of the page. Show the balance before you work out any interest.', 0.0)
ws.doc.add_paragraph()

ws.Q(T('A camera has a cash price of $1200. A deposit of 15% is paid and the balance is repaid '
       'over 2 years at 10% per annum simple interest.'))
ws.SQ(T('Find the deposit.'), marks=1)
ws.SQ(T('Find the monthly instalment.'), marks=3)

ws.Q(T('A bicycle has a cash price of $960. Under hire purchase a deposit of $160 is paid, '
       'followed by 24 monthly instalments of $38. Find the total hire purchase price, and how '
       'much more than the cash price is paid.'), marks=3)

ws.Q(T('A piano is bought for a deposit of $2000 and 60 monthly instalments of $130. The balance '
       'was charged at 6% per annum simple interest. Find the cash price of the piano.'), marks=4)

ws.Q(T('A refrigerator has a cash price of $1500. A customer pays a deposit of 20% and repays the '
       'balance over 2 years in monthly instalments of $55. Find the rate of simple interest per '
       'annum.'), marks=4)

ws.Q(T('A television has a cash price of $1400. A deposit of $200 is paid and the balance is '
       'charged at 10% per annum simple interest, repaid in monthly instalments of $60. Find the '
       'number of years of the agreement.'), marks=4)

ws.Q(T('A tablet has a cash price of $600. Plan A is a deposit of 20% with the balance over '
       '1 year at 12% per annum. Plan B has no deposit, with the full amount over 2 years at 9% '
       'per annum. Determine which plan costs less in total.'), marks=5)

ws.doc.add_paragraph()
mono_box(ws, 'Answers', [
 "1.  (a) $180                (b) $51",
 "2.  Hire purchase price $1072,  $112 more than the cash price",
 "3.  Balance $6000,  cash price $8000",
 "4.  5% per annum",
 "5.  2 years",
 "6.  Plan A $657.60,  Plan B $708  →  Plan A costs $50.40 less",
 "    (Plan B has the smaller instalment: $29.50 against $44.80)",
])

ws.save('hire_purchase_worked_examples.docx')
print('saved')
