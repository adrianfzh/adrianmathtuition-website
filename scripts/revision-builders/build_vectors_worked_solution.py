# -*- coding: utf-8 -*-
"""Vectors — worked solution for the OAB / BC:CA = 2:1 question.

Three rendering facts drove how this is built, all established by probe:

  * pandoc emits ONE OMML run per character, so "OAB" is three runs and both
    Word and LibreOffice space them as three italic variables — `O A B`.
    ommlfix.merge_runs folds adjacent same-format runs back together, which
    fixes point names and puts the arrow of \\overrightarrow{OA} over a tight
    "OA" instead of over a gap.
  * \\mathbf{a} DOES carry <m:sty m:val="b"> in the XML — Word honours it,
    LibreOffice ignores it. Since the delivered PDF is rendered through
    LibreOffice, the bold would silently vanish from it, so every vector
    letter is a bold Word run instead of math. Rule: no math run may contain
    a vector letter.
  * A superscript written as its own equation renders as a placeholder box,
    so cm² uses the Unicode character.

Vulgar fractions are Unicode where one exists (½ ⅔ ¼ ¾) and real math only
where it does not, which keeps bold letters and fractions from ever needing
to share a run.
"""
import worksheet_lib
from worksheet_lib import Worksheet
from ommlfix import merge_runs
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Every equation in this document goes through the run merger.
_raw_omml = worksheet_lib._latex_to_omml
worksheet_lib._latex_to_omml = lambda expr, display=False: (
    lambda el: merge_runs(el) if el is not None else None)(_raw_omml(expr, display=display))

NAVY = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x59, 0x59, 0x59)
BAR = 'DEEBF7'
EDGE = '2E5C8A'
REASON_TAB = 10.4          # cm — the right-hand "why" column

ws = Worksheet()
ws.title('O-Level Mathematics')
ws.subtitle('Vectors — Worked Solution')


# ---------------------------------------------------------------- box chrome
def _shade(cell, fill):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:color'), 'auto')
    el.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(el)


def _box(table, color):
    b = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '6')
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)
        b.append(e)
    tblPr = table._tbl.tblPr
    anchor = None
    for tag in ('w:tblLayout', 'w:tblCellMar', 'w:tblLook'):
        anchor = tblPr.find(qn(tag))
        if anchor is not None:
            break
    (anchor.addprevious(b) if anchor is not None else tblPr.append(b))


# ---------------------------------------------------------------- atoms
def t(s):
    return ('text', s)


def m(s):
    return ('math', s)


def bold(s):
    return ('text', s, {'bold': True})


def P(s):
    """A point or segment name — italic, e.g. OAB, BC, APQ."""
    return ('text', s, {'italic': True})


def V(name):
    r"""A vector with an arrow over it: \overrightarrow{OA}, rendered tight."""
    return m(r'\overrightarrow{' + name + '}')


def E(s):
    """Algebra with bold vector letters marked by @: '6@b + ⅔(3@a)'.

    Any other letter is a scalar variable (m, n) and is set italic, so an m
    inside an expression matches the italic m in the surrounding prose rather
    than coming out upright.
    """
    parts, buf, i = [], '', 0

    def flush():
        nonlocal buf
        if buf:
            parts.append(t(buf))
            buf = ''

    while i < len(s):
        c = s[i]
        if c == '@' and i + 1 < len(s):
            flush()
            parts.append(bold(s[i + 1]))
            i += 2
        elif c.isalpha():
            flush()
            parts.append(P(c))
            i += 1
        else:
            buf += c
            i += 1
    flush()
    return parts


# ---------------------------------------------------------------- structure
def question(parts, indent=0.0, space_before=4):
    p = ws.doc.add_paragraph()
    p.style = ws.doc.styles['Normal']
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(space_before)
    ws._fill(p, parts)
    return p


def solution(title, steps):
    """One navy box: working on the left, the reason for the step on the right.

    A step is (indent, parts) or (indent, parts, reason); the reason sits at a
    fixed tab so the right-hand column lines up all the way down the box.
    """
    tb = ws.doc.add_table(rows=2, cols=1)
    tb.autofit = False
    _box(tb, EDGE)

    head = tb.rows[0].cells[0]
    head.width = Cm(16)
    _shade(head, BAR)
    hp = head.paragraphs[0]
    hp.paragraph_format.line_spacing = 1.0
    hp.paragraph_format.space_before = Pt(2)
    hp.paragraph_format.space_after = Pt(2)
    hr = hp.add_run(title)
    hr.font.name = 'Calibri'
    hr.font.size = Pt(10.5)
    hr.bold = True
    hr.font.color.rgb = NAVY

    body = tb.rows[1].cells[0]
    body.width = Cm(16)
    first = True
    for step in steps:
        indent, parts = step[0], step[1]
        reason = step[2] if len(step) > 2 else None
        bp = body.paragraphs[0] if first else body.add_paragraph()
        first = False
        pf = bp.paragraph_format
        pf.line_spacing = 1.0
        pf.space_before = Pt(3 if indent == 0 else 1)
        pf.space_after = Pt(0)
        pf.left_indent = Cm(indent)
        if reason:
            pf.tab_stops.add_tab_stop(Cm(max(REASON_TAB - indent, 1.0)))
        ws._fill(bp, parts)
        if reason:
            r = bp.add_run('\t' + reason)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(9)
            r.italic = True
            r.font.color.rgb = GREY
    ws.doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return tb


def remark(parts):
    p = ws.doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(9)
    for part in parts:
        kind = part[0]
        if kind == 'text':
            attrs = dict(part[2]) if len(part) > 2 else {}
            r = p.add_run(part[1])
            r.font.name = 'Times New Roman'
            r.font.size = Pt(9)
            r.italic = not attrs.get('bold')
            r.bold = attrs.get('bold', False)
            r.font.color.rgb = GREY
        else:
            ws._fill(p, [part])
    return p


# ================================================================ the question
ws.figure('vec_fig.png', width_cm=10.5)

question([P('OAB'), t(' is a triangle. '), P('C'), t(' is the point on '),
          P('BA'), t(' such that '), P('BC'), t(' : '), P('CA'), t(' = 2 : 1.   '),
          V('OA'), t(' = ')] + E('3@a') + [t(' and '), V('OB'), t(' = ')] + E('6@b')
         + [t('.')], space_before=8)

question([bold('(a)'), t('  Show that the position vector of '), P('C'),
          t(' is given by '), V('OC'), t(' = ')] + E('2@a + 2@b') + [t('.')],
         space_before=8)

question([bold('(b)'), t('  '), P('P'), t(' is the midpoint of '), P('OC'),
          t(' and '), P('Q'), t(' is a point on '), P('OB'), t(' such that '),
          P('APQ'), t(' is a straight line.')])
question([V('AQ'), t(' = '), P('m'), t(' '), V('AP'), t('   and   '), V('OQ'),
          t(' = '), P('n'), t(' '), V('OB'), t('   where '), P('m'), t(' and '),
          P('n'), t(' are numbers. Find the ratio '), P('OQ'), t(' : '),
          P('QB'), t('.')], indent=1.0, space_before=2)

question([bold('(c)'), t('  The area of triangle '), P('OBC'),
          t(' is 25 cm². Find the area of triangle '), P('OAC'), t('.')],
         space_before=6)

ws.page_break()

# ================================================================== part (a)
# C is two thirds of the way from B to A, so build OC by going O -> B -> C.
solution('Solution — (a)', [
    (0.0, [t('Since '), P('BC'), t(' : '), P('CA'), t(' = 2 : 1, '), P('C'),
           t(' divides '), P('BA'), t(' so that '), V('BC'), t(' = ⅔ '),
           V('BA'), t('.')]),
    (0.0, [V('BA'), t(' = '), V('OA'), t(' − '), V('OB'), t(' = ')]
          + E('3@a − 6@b'), 'from B towards A'),
    (0.0, [V('OC'), t(' = '), V('OB'), t(' + '), V('BC')], 'O to B, then B to C'),
    (1.3, [t('= ')] + E('6@b + ⅔(3@a − 6@b)')),
    (1.3, [t('= ')] + E('6@b + 2@a − 4@b')),
    (1.3, [t('= ')] + E('2@a + 2@b') + [t('     (shown)')]),
])
remark([t('Going O → B → C keeps every step along a vector you already '
          'know. O → A → C works just as well: '), V('AC'), t(' = ⅓ '),
        V('AB'), t(' = ')] + E('2@b − @a') + [t(', so '), V('OC'), t(' = ')]
       + E('3@a + 2@b − @a') + [t(' = ')] + E('2@a + 2@b') + [t('.')])

# ================================================================== part (b)
# The method: write OQ two ways, then compare coefficients — legal only
# because a and b are non-parallel.
solution('Solution — (b)', [
    (0.0, [V('OP'), t(' = ½ '), V('OC'), t(' = ½(')] + E('2@a + 2@b')
          + [t(') = ')] + E('@a + @b'), 'P is the midpoint of OC'),
    (0.0, [V('AP'), t(' = '), V('OP'), t(' − '), V('OA'), t(' = (')]
          + E('@a + @b') + [t(') − ')] + E('3@a') + [t(' = ')]
          + E('@b − 2@a')),
    (0.0, [t('Along the straight line '), P('APQ'), t(':')]),
    (1.3, [V('AQ'), t(' = '), P('m'), t(' '), V('AP'), t(' = '), P('m'),
           t('(')] + E('@b − 2@a') + [t(') = ')] + E('−2m@a + m@b')),
    (1.3, [V('OQ'), t(' = '), V('OA'), t(' + '), V('AQ'), t(' = ')]
          + E('(3 − 2m)@a + m@b') + [t('        … (1)')]),
    (0.0, [t('Since '), P('Q'), t(' lies on '), P('OB'), t(':')]),
    (1.3, [V('OQ'), t(' = '), P('n'), t(' '), V('OB'), t(' = ')] + E('6n@b')
          + [t('        … (2)')]),
    (0.0, [bold('a'), t(' and '), bold('b'),
           t(' are non-parallel, so compare coefficients in (1) and (2):')]),
    (1.3, [bold('a'), t(' :   3 − 2'), P('m'), t(' = 0     so     '),
           m(r'm=\dfrac{3}{2}')]),
    (1.3, [bold('b'), t(' :   '), P('m'), t(' = 6'), P('n'), t('     so     '),
           m(r'n=\dfrac{1}{6}\times\dfrac{3}{2}=\dfrac{1}{4}')]),
    (0.0, [V('OQ'), t(' = ¼ '), V('OB'), t('     and     '), V('QB'),
           t(' = '), V('OB'), t(' − '), V('OQ'), t(' = ¾ '), V('OB')]),
    (0.0, [bold('OQ : QB = ¼ : ¾ = 1 : 3')]),
])
remark([t('Comparing coefficients is the whole method, and it is only valid '
          'because '), bold('a'), t(' and '), bold('b'),
        t(' are non-parallel — that is what forces the two expressions for '),
        V('OQ'), t(' to match term by term. Write '), V('OQ'),
        t(' two different ways, then equate.')])

# ================================================================== part (c)
# Same apex, bases on the same line, so equal heights: the areas carry the
# base ratio. No vectors needed at all.
solution('Solution — (c)', [
    (0.0, [t('Triangles '), P('OBC'), t(' and '), P('OAC'),
           t(' have their bases '), P('BC'), t(' and '), P('CA'),
           t(' on the same line '), P('BA'), t(', and share the apex '),
           P('O'), t('.')]),
    (0.0, [t('So they have the same perpendicular height, and their areas are '
             'in the ratio of their bases:')]),
    (1.3, [m(r'\dfrac{\text{area of } OBC}{\text{area of } OAC}'
             r'=\dfrac{BC}{CA}=\dfrac{2}{1}')]),
    (0.0, [t('area of '), P('OAC'), t(' = ½ × 25 = '),
           bold('12.5 cm²')]),
])
remark([t('No vectors are needed here. Two triangles with the same height have '
          'areas in the ratio of their bases — and '), P('BC'), t(' : '),
        P('CA'), t(' = 2 : 1 was given in the very first line.')])

ws.save('vectors_worked_solution.docx')
print('saved')
