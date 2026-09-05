#!/usr/bin/env python3
"""The six S1 revision worksheets."""
from batch_lib import build
from build_lib import T, B, I, M
import revision_lib as R

LL = 'Sec 1 Mathematics Revision'


def _mistakes(ws, items):
    ws.para([B('Mistakes to avoid')])
    for i, e in enumerate(items, 1):
        ws.para([T(f'{i}.  ')] + R.split_math(e))


def _gradient_table(ws):
    """The four gradients as pictures, side by side (Adrian, 5 Sep 2026 — "put in
    diagrams ... perhaps in table form for visual explanation"). Sketches are
    committed PNGs; regenerate with scripts/revision-builders/make_gradient_sketches.py."""
    from pathlib import Path
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    assets = Path(__file__).resolve().parent / "assets"
    cols = [
        ("grad-positive",  "Positive",  "m > 0",       "rises left to right"),
        ("grad-negative",  "Negative",  "m < 0",       "falls left to right"),
        ("grad-zero",      "Zero",      "m = 0",       "horizontal, y = c"),
        ("grad-undefined", "Undefined", "no m",        "vertical, x = k"),
    ]
    t = ws.doc.add_table(rows=4, cols=len(cols))
    t.style = ws.doc.styles["Table Grid"]
    t.autofit = False
    for r, row in enumerate(t.rows):
        for c, cell in enumerate(row.cells):
            cell.width = Cm(3.9)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = para.paragraph_format.space_after = None
            name, label, sym, note = cols[c]
            if r == 0:
                para.add_run().add_picture(str(assets / f"{name}.png"), width=Cm(2.7))
            elif r == 1:
                ws._fill(para, [B(label)])
            elif r == 2:
                ws._fill(para, [M(sym)])
            else:
                ws._fill(para, R.split_math(note))
    ws.doc.add_paragraph()
    return t


def n_primes(ws):
    ws.para([B('Notes:')])
    ws.para([B('Prime factorisation')])
    ws.para([T('A prime number has exactly two factors, 1 and itself. 1 is '), B('not'),
             T(' prime; 2 is the only even prime.')])
    ws.para([T('Write every number as a product of primes in index form:')])
    ws.math_block(r'360 = 2^3 \times 3^2 \times 5')
    ws.para([B('HCF and LCM from the index form')])
    ws.para([T('HCF — take each '), B('common'), T(' prime to its '), B('lowest'), T(' power.')])
    ws.para([T('LCM — take '), B('every'), T(' prime that appears, to its '), B('highest'),
             T(' power.')])
    ws.math_block(r'a = 2^3 \times 3^2, \quad b = 2^2 \times 3 \times 5')
    ws.math_block(r'\text{HCF} = 2^2 \times 3 = 12 \qquad \text{LCM} = 2^3 \times 3^2 \times 5 = 360')
    ws.para([B('Squares and cubes')])
    ws.para([T('A number is a perfect '), B('square'),
             T(' exactly when every index in its prime factorisation is even; a perfect '),
             B('cube'), T(' when every index is a multiple of 3.')])
    ws.para([T('To make a number a square, multiply by whatever the odd indices are missing.')])
    ws.math_block(r'\sqrt{2^4 \times 3^6} = 2^2 \times 3^3 \qquad \sqrt[3]{2^6 \times 5^3} = 2^2 \times 5')
    ws.para([B('Which one does the word problem want?')])
    ws.para([T('Splitting into equal groups, or the largest tile that fits — '), B('HCF'), T('.')])
    ws.para([T('Events repeating together, or the smallest length made from both — '), B('LCM'), T('.')])
    _mistakes(ws, [
        'Calling 1 a prime number.',
        'Taking the highest power for the HCF and the lowest for the LCM.',
        'Including a prime in the HCF that only one number has.',
        'Choosing LCM when the question is about sharing equally.',
    ])


def n_algebra(ws):
    ws.para([B('Notes:')])
    ws.para([B('Writing expressions')])
    ws.para([T('"5 more than '), M('x'), T('" is '), M('x + 5'), T('. "5 less than '), M('x'),
             T('" is '), M('x - 5'), T(', not '), M('5 - x'), T(' — order matters when subtracting.')])
    ws.para([B('Like terms')])
    ws.para([T('Only terms with exactly the same letters and powers can be added.')])
    ws.math_block(r'3x + 5x = 8x \qquad 3x + 5y \text{ stays as it is} \qquad 3x + 3x^2 \text{ does not combine}')
    ws.para([B('Expanding')])
    ws.math_block(r'a(b + c) = ab + ac')
    ws.para([T('A minus in front of a bracket changes '), B('every'), T(' sign inside:')])
    ws.math_block(r'-2(x - 4) = -2x + 8')
    ws.para([B('Factorising')])
    ws.para([T('Take out the highest common factor of the numbers '), I('and'), T(' the letters:')])
    ws.math_block(r'12x^2 + 18x = 6x(2x + 3)')
    ws.para([B('Solving a linear equation')])
    ws.para([T('Do the same thing to both sides. Move the letters to one side and the numbers '
               'to the other, then divide.')])
    ws.para([T('If there are fractions, multiply '), B('every term'), T(' by the denominator first.')])
    ws.para([B('Substitution')])
    ws.para([T('Use brackets for negative values: if '), M('x = -2'), T(' then '),
             M('3x^2 = 3(-2)^2 = 12'), T(', not '), M('-12'), T('.')])
    _mistakes(ws, [
        'Writing "5 less than $x$" as $5 - x$.',
        'Adding unlike terms such as $2x$ and $3x^2$.',
        'Changing only the first sign after a minus in front of a bracket.',
        'Multiplying only some terms when clearing a fraction from an equation.',
    ])


def n_rrps(ws):
    ws.para([B('Notes:')])
    ws.para([B('Ratio')])
    ws.para([T('A ratio compares parts. Simplify by dividing by the HCF, and make the units the '
               'same before you simplify.')])
    ws.para([T('To share an amount in the ratio '), M('a : b'), T(', there are '), M('a + b'),
             T(' equal parts; find the value of one part first.')])
    ws.para([B('Rate')])
    ws.para([T('A rate compares two '), I('different'), T(' quantities — dollars per kg, litres '
               'per minute. Always write the unit with the number.')])
    ws.para([B('Percentage')])
    ws.math_block(r'\text{percentage} = \dfrac{\text{part}}{\text{whole}} \times 100\%')
    ws.para([T('Increase by '), M('r\\%'), T(': multiply by '), M(r'\left(1 + \tfrac{r}{100}\right)'),
             T('.   Decrease: multiply by '), M(r'\left(1 - \tfrac{r}{100}\right)'), T('.')])
    ws.para([T('Percentage change is always measured against the '), B('original'), T(' amount:')])
    ws.math_block(r'\text{percentage change} = \dfrac{\text{new} - \text{original}}{\text{original}} \times 100\%')
    ws.para([B('Reverse percentage')])
    # Currency written through the math renderer, not as bare "$60" in a text
    # run: a plain dollar sign is the inline-math delimiter everywhere else in
    # this pipeline, and two of them in one paragraph is a trap waiting to fire.
    ws.para([T('If a price '), I('after'), T(' a 20% increase is '), M(r'\$60'), T(', then '),
             M(r'\$60'), T(' is 120% of the original — divide, do not take 20% off.')])
    ws.math_block(r'\text{original} = \dfrac{60}{1.2} = \$50')
    ws.para([B('Speed')])
    ws.math_block(r'\text{speed} = \dfrac{\text{distance}}{\text{time}} \qquad \text{distance} = \text{speed} \times \text{time} \qquad \text{time} = \dfrac{\text{distance}}{\text{speed}}')
    ws.para([T('Average speed is total distance over total time — never the average of the speeds.')])
    ws.para([T('To convert km/h to m/s, multiply by '), M(r'\tfrac{1000}{3600} = \tfrac{5}{18}'), T('.')])
    ws.para([T('Write 2 h 30 min as 2.5 h, not 2.30 h.')])
    _mistakes(ws, [
        'Measuring a percentage change against the new amount instead of the original.',
        'Taking 20% off to reverse a 20% increase.',
        'Averaging two speeds instead of dividing total distance by total time.',
        'Using 2.30 for 2 hours 30 minutes.',
    ])


def n_linear(ws):
    ws.para([B('Notes:')])
    ws.para([B('The equation of a straight line')])
    ws.math_block(r'y = mx + c')
    ws.para([M('m'), T(' is the gradient and '), M('c'), T(' is the '), M('y'), T('-intercept, '
             'the value of '), M('y'), T(' where the line crosses the '), M('y'), T('-axis.')])
    ws.para([B('Gradient')])
    ws.math_block(r'm = \dfrac{\text{rise}}{\text{run}} = \dfrac{y_2 - y_1}{x_2 - x_1}')
    ws.para([T('Which way the line leans is the whole of '), M('m'),
             T(' — read the sign off the picture before you read anything else.')])
    _gradient_table(ws)
    ws.para([B('Reading a graph')])
    ws.para([T('Pick two points that sit '), I('exactly'), T(' on grid intersections, as far '
             'apart as you can, and use them for the gradient. Two close points make a small '
             'reading error large.')])
    ws.para([B('Finding the equation from information')])
    ws.para([T('1.  Find '), M('m'), T(' from two points, or from a parallel line — parallel '
             'lines have equal gradients.')])
    ws.para([T('2.  Substitute one known point into '), M('y = mx + c'), T(' and solve for '),
             M('c'), T('.')])
    ws.para([B('Is a point on the line?')])
    ws.para([T('Substitute its coordinates. If both sides balance, it lies on the line.')])
    ws.para([B('Intercepts')])
    ws.para([T('For the '), M('y'), T('-intercept put '), M('x = 0'), T('; for the '), M('x'),
             T('-intercept put '), M('y = 0'), T('.')])
    _mistakes(ws, [
        'Subtracting the coordinates in a different order on the top and the bottom of the gradient.',
        'Reading the gradient off the graph without checking the scales on the two axes.',
        'Confusing the $x$- and $y$-intercepts.',
        'Giving a vertical line the equation $y = k$; it is $x = k$, and it has no gradient.',
    ])


def n_mensuration(ws):
    ws.para([B('Notes:')])
    ws.para([B('Area and perimeter')])
    ws.math_block(r'\text{rectangle} = l \times b \qquad \text{triangle} = \tfrac{1}{2} \times b \times h')
    ws.math_block(r'\text{parallelogram} = b \times h \qquad \text{trapezium} = \tfrac{1}{2}(a + b) \times h')
    ws.para([T('The height is always '), B('perpendicular'), T(' to the base — not the slanted side.')])
    ws.para([B('Circles')])
    ws.math_block(r'\text{circumference} = 2\pi r \qquad \text{area} = \pi r^2')
    ws.para([T('Read carefully whether you are given the radius or the diameter.')])
    ws.para([B('Volume and surface area')])
    ws.math_block(r'\text{prism} = \text{area of cross-section} \times \text{length}')
    ws.math_block(r'\text{cylinder: } V = \pi r^2 h \qquad \text{curved surface} = 2\pi r h')
    ws.para([T('A closed cylinder has two circular ends as well: '),
             M(r'\text{total} = 2\pi r h + 2\pi r^2'), T('.')])
    ws.para([B('Composite figures')])
    ws.para([T('Split the shape into pieces you know, work each out separately, then add — or '
               'subtract, if a piece has been cut away. Draw the split on the diagram.')])
    ws.para([B('Units')])
    ws.para([T('Length in cm, area in cm'), M('^2'), T(', volume in cm'), M('^3'),
             T('. Convert before calculating, not after.')])
    ws.math_block(r'1 \text{ m}^2 = 10\,000 \text{ cm}^2 \qquad 1 \text{ m}^3 = 1\,000\,000 \text{ cm}^3 \qquad 1 \text{ litre} = 1000 \text{ cm}^3')
    _mistakes(ws, [
        'Using a slanted side as the height of a triangle or trapezium.',
        'Using the diameter where the formula asks for the radius.',
        'Forgetting the two ends when finding the total surface area of a closed cylinder.',
        'Mixing cm and m within one calculation.',
    ])


def n_stats(ws):
    ws.para([B('Notes:')])
    ws.para([B('Collecting and showing data')])
    ws.para([T('Pictogram — one symbol stands for a number of items; always read the key.')])
    ws.para([T('Bar graph — the height is the frequency; the bars are separate and equally wide.')])
    ws.para([T('Pie chart — the whole circle is the total.')])
    ws.math_block(r'\text{angle of a sector} = \dfrac{\text{frequency}}{\text{total}} \times 360^\circ')
    ws.para([T('Line graph — for data that changes over time.')])
    ws.para([B('The three averages')])
    ws.para([T('Mean '), M(r'= \dfrac{\text{sum of the values}}{\text{how many values}}'), T('.')])
    ws.para([T('Median — the middle value once the data is in '), B('order'), T('. With an even '
             'number of values, take the mean of the middle two.')])
    ws.para([T('Mode — the most frequent value.')])
    ws.para([B('Choosing which average to quote')])
    ws.para([T('The mean uses every value but is dragged by an extreme one. The median ignores '
               'extremes. The mode is the only average that works for non-numerical data such '
               'as favourite colour.')])
    ws.para([B('From a frequency table')])
    ws.para([T('The mean is '), M(r'\dfrac{\sum fx}{\sum f}'), T(' — multiply each value by its '
             'frequency, add, then divide by the total frequency. Do not divide by the number '
             'of rows.')])
    _mistakes(ws, [
        'Finding the median without ordering the data.',
        'Dividing by the number of rows instead of the total frequency.',
        'Ignoring the key on a pictogram.',
        'Reading a pie chart sector as a frequency rather than a fraction of the total.',
    ])


SHEETS = [
    dict(level='S1', folder='S1', level_line=LL,
         tags=['Numbers (Prime Factorization)', 'Numbers (HCF and LCM)'],
         title='Prime Factorisation, HCF and LCM',
         filename='01 Prime Factorisation HCF and LCM Revision.docx', notes=n_primes,
         skill_titles=['Prime Factorisation in Index Form', 'Choosing Between HCF and LCM']),
    dict(level='S1', folder='S1', level_line=LL,
         tags=['Algebra (Linear Equations)', 'Algebra (Expressions)', 'Algebra (Factorization)'],
         title='Algebra', filename='06 Algebra Revision.docx', notes=n_algebra,
         skill_titles=['Solving a Linear Equation', 'Forming an Expression From Words']),
    dict(level='S1', folder='S1', level_line=LL,
         tags=['Numbers (Rate)', 'Numbers (Ratio)', 'Numbers (Percentages)', 'Numbers (Speed)'],
         title='Rate, Ratio, Percentages and Speed',
         filename='12 Rate Ratio Percentages and Speed Revision.docx', notes=n_rrps,
         skill_titles=['Sharing in a Given Ratio', 'Percentage Change and Reverse Percentage']),
    dict(level='S1', folder='S1', level_line=LL, tags=['Coordinate Geometry (Lines)'],
         title='Linear Functions', filename='19 Linear Functions Revision.docx', notes=n_linear,
         skill_titles=['Finding the Gradient and the y-intercept',
                       'Finding the Equation of a Line']),
    dict(level='S1', folder='S1', level_line=LL, tags=['Mensuration'],
         title='Mensuration', filename='20 Mensuration Revision.docx', notes=n_mensuration,
         skill_titles=['Area and Perimeter of a Composite Figure',
                       'Volume and Surface Area of a Prism or Cylinder']),
    dict(level='S1', folder='S1', level_line=LL, tags=['Statistics'],
         title='Statistics', filename='21 Statistics Revision.docx', notes=n_stats,
         skill_titles=['Reading a Statistical Diagram', 'Mean, Median and Mode']),
]

if __name__ == "__main__":
    for cfg in SHEETS:
        print(f"\n=== {cfg['title']}")
        try:
            build(cfg)
        except Exception as e:
            print(f"   !! FAILED: {type(e).__name__}: {e}")
