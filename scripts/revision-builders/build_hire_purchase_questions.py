# -*- coding: utf-8 -*-
"""Hire purchase — six exam questions, one per calculation type.

Chosen so no two questions are solved the same way; between them they cover
every direction the topic is asked in:

    1  instalment      cash price + HP price + deposit %      -> monthly payment
    2  deposit         cash price + HP price + instalments    -> deposit as a %
    3  interest rate   cash price + deposit % + instalments   -> rate p.a.
    4  forward, full   price + down payment % + rate + term   -> deposit, interest, instalment
    5  reverse algebra down payment % + rate + term + instal. -> cash price
    6  comparison      two plans against one cash price       -> which is cheaper

Numbering is written literally rather than through the library's autonumber so
a one-line question rides up onto the number line ("1.  The cash price ...").
Q4 has a stem and three parts, so its number line carries the stem and the
parts follow at the label column.

Answers sit directly under each question, no working space.  Provenance stays
off the sheet.
"""
from worksheet_lib import Worksheet
from docx.shared import Cm, Pt
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.ns import qn

NUM_COL, LABEL_COL, TEXT_COL, MARK_COL = 0.0, 1.0, 2.0, 15.5

ws = Worksheet()
ws.title('O-Level Mathematics')
ws.subtitle('Hire Purchase')


def _run(p, text, bold=False):
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9.5)
    r.bold = bold
    for wt in r._r.findall(qn('w:t')):
        wt.set(qn('xml:space'), 'preserve')
    return r


def _marks(p, marks):
    if marks is None:
        return
    p.paragraph_format.tab_stops.add_tab_stop(Cm(MARK_COL), WD_TAB_ALIGNMENT.RIGHT)
    _run(p, f'\t[{marks}]')


def _hang(space_before=6):
    p = ws.doc.add_paragraph()
    p.style = ws.doc.styles['Normal']
    pf = p.paragraph_format
    pf.left_indent = Cm(TEXT_COL)
    pf.first_line_indent = Cm(-TEXT_COL)
    pf.space_before = Pt(space_before)
    pf.tab_stops.add_tab_stop(Cm(LABEL_COL))
    pf.tab_stops.add_tab_stop(Cm(TEXT_COL))
    return p


def Qsingle(num, parts, marks=None):
    """A question with no parts: `1.  The cash price ...`."""
    p = _hang()
    _run(p, f'{num}.', bold=True)
    _run(p, '\t\t')
    ws._fill(p, parts)
    _marks(p, marks)
    return p


def Qstem(num, parts):
    """A stem that its (a),(b),(c) will hang beneath."""
    p = _hang()
    _run(p, f'{num}.', bold=True)
    _run(p, '\t\t')
    ws._fill(p, parts)
    return p


def Pline(label, parts, marks=None):
    """A lettered part, aligned to the label and text columns."""
    p = ws.doc.add_paragraph()
    p.style = ws.doc.styles['Normal']
    pf = p.paragraph_format
    pf.left_indent = Cm(TEXT_COL)
    pf.first_line_indent = Cm(-(TEXT_COL - LABEL_COL))
    pf.tab_stops.add_tab_stop(Cm(TEXT_COL))
    _run(p, f'({label})')
    _run(p, '\t')
    ws._fill(p, parts)
    _marks(p, marks)
    return p


def Sub(parts, space_before=2):
    """An indented line inside a question — a plan, an option, a table row."""
    p = ws.doc.add_paragraph()
    p.style = ws.doc.styles['Normal']
    pf = p.paragraph_format
    pf.left_indent = Cm(TEXT_COL)
    pf.space_before = Pt(space_before)
    ws._fill(p, parts)
    return p


def note(text):
    p = ws.doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9)
    r.italic = True
    p.paragraph_format.space_after = Pt(6)
    return p


def t(s):
    return ('text', s)


def m(s):
    return ('math', s)


def b(s):
    return ('text', s, {'bold': True})


note('Interest on a hire purchase is simple interest, and it is charged on '
     'the amount still owing after the deposit — never on the full cash price.')

# ------------------------------------------------- 1  find the instalment
# HP price - deposit, shared over the payments.
Qsingle(1, [t('The cash price of a refrigerator is $2499. The hire-purchase '
              'price of the refrigerator is $2999. The hire-purchase price is '
              'a deposit of 25% of the cash price plus 24 equal monthly '
              'payments. Calculate one monthly payment.')], marks=2)
ws.ans([t('$98.93')])

# ------------------------------------------------- 2  find the deposit, as a %
# The deposit is what the instalments do not account for.
Qsingle(2, [t('The cash price of a television is $1750. The hire-purchase '
              'price of the television is $2005. The hire-purchase price is a '
              'deposit plus 12 monthly payments of $135. Calculate the deposit '
              'as a percentage of the cash price.')], marks=4)
ws.ans([t('22%')])

# ------------------------------------------------- 3  find the rate p.a.
# Interest = total paid - cash price; principal = cash price - deposit.
Qsingle(3, [t('The cash price of a washing machine is $840. Under a '
              'hire-purchase scheme the deposit is 15% of the cash price and '
              'the 24 subsequent equal monthly payments are $33.50 each. '
              'Calculate the interest rate per annum.')], marks=3)
ws.ans([t('6.30% (3 s.f.)')])

# ------------------------------------------------- 4  the full forward chain
Qstem(4, [t('Mary bought a washing machine for $2840. She paid a down payment '
            'of 20% and repaid the balance in 36 equal monthly instalments. '
            'The interest charged was 5% per annum on the remaining amount. '
            'Find')])
Pline('a', [t('the amount of the down payment,')], marks=2)
Pline('b', [t('the interest paid,')], marks=2)
Pline('c', [t('the amount of each monthly instalment.')], marks=2)
ws.ans([t('(a) $568   (b) $340.80   (c) $72.58')])

# ------------------------------------------------- 5  reverse: find the price
# Everything is a multiple of x, so the whole payment collapses to 0.88x.
Qsingle(5, [t('The price of a sofa bed is $'), m('x'),
            t('. Toby buys it on hire purchase. He pays a down payment of 25% '
              'and arranges to pay the remaining amount in monthly instalments '
              'over 26 months, at a simple interest rate of 8% per annum. '
              'Given that his monthly instalment is $88, find '), m('x'),
            t('.')], marks=4)
ws.ans([('text', 'x', {'italic': True}), t(' = 2600')])

# ------------------------------------------------- 6  compare two plans
Qstem(6, [t('Jered wants to buy a handphone priced at $1299. He may pay for it '
            'using either of the following packages.')])
Sub([b('Package A:'), t('  15% deposit and 12 monthly payments of $95.')],
    space_before=4)
Sub([b('Package B:'), t('  $400 upfront and 24 monthly payments of $38.50.')])
p = Sub([t('Determine which package is cheaper, and by how much.')], space_before=4)
_marks(p, 3)
ws.ans([t('Package B, by $10.85')])

ws.save('hire_purchase_questions.docx')
print('saved')
