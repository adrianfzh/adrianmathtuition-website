# -*- coding: utf-8 -*-
"""H2 inequalities — the five hardest-to-spot substitutions.

Five "hence solve" questions where the second inequality is the first one
under a replacement that is deliberately hard to see:

    1  x -> 1/x        hidden behind an x^2 denominator
    2  x -> -x         plus a multiply-by-(-1) that reverses the inequality
    3  x -> -|x|       a compound: plain |x| fails
    4  x -> sin x      only after rewriting cosec x and cos^2 x
    5  x -> x/2        a scaling, and the answer is the complement

Numbering is written literally rather than through the library's autonumber
so the first part can ride up onto the number line — "1.  (a)  Solve ..." —
instead of leaving the number alone on a line of its own.  Everything after a
hoist must therefore use literal labels too; the library's SQ() would restart
its own (a),(b) sequence and fight this one.

Answers sit directly under each question and name the replacement, since the
replacement is the thing being taught.  Provenance stays off the sheet.
"""
from worksheet_lib import Worksheet
from docx.shared import Cm, Pt
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.ns import qn

NUM_COL, LABEL_COL, TEXT_COL, MARK_COL = 0.0, 1.0, 2.0, 15.5

ws = Worksheet()
ws.title('H2 Mathematics')
ws.subtitle('Inequalities — Solving a Related Inequality by Substitution')


def _run(p, text, bold=False):
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9.5)
    r.bold = bold
    for wt in r._r.findall(qn('w:t')):
        wt.set(qn('xml:space'), 'preserve')
    return r


def _marks(p, marks):
    if marks is None:
        return
    p.paragraph_format.tab_stops.add_tab_stop(Cm(MARK_COL), WD_TAB_ALIGNMENT.RIGHT)
    _run(p, f'\t[{marks}]')


def Qline(num, label, parts, marks=None):
    """First part hoisted onto the number line: `1.  (a)  Solve ...`."""
    p = ws.doc.add_paragraph()
    p.style = ws.doc.styles['Normal']
    pf = p.paragraph_format
    pf.left_indent = Cm(TEXT_COL)
    pf.first_line_indent = Cm(-TEXT_COL)
    pf.space_before = Pt(6)
    pf.tab_stops.add_tab_stop(Cm(LABEL_COL))
    pf.tab_stops.add_tab_stop(Cm(TEXT_COL))
    _run(p, f'{num}.', bold=True)
    _run(p, '\t')
    _run(p, f'({label})')
    _run(p, '\t')
    ws._fill(p, parts)
    _marks(p, marks)
    return p


def Pline(label, parts, marks=None):
    """A later part, aligned to the same label and text columns."""
    p = ws.doc.add_paragraph()
    p.style = ws.doc.styles['Normal']
    pf = p.paragraph_format
    pf.left_indent = Cm(TEXT_COL)
    pf.first_line_indent = Cm(-(TEXT_COL - LABEL_COL))
    pf.tab_stops.add_tab_stop(Cm(TEXT_COL))
    _run(p, f'({label})')
    _run(p, '\t')
    ws._fill(p, parts)
    _marks(p, marks)
    return p


def note(text):
    p = ws.doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9)
    r.italic = True
    p.paragraph_format.space_after = Pt(6)
    return p


def t(s):
    return ('text', s)


def m(s):
    return ('math', s)


note('Do not use a calculator in any question. In each one the second '
     'inequality is the first one after a substitution — find the '
     'replacement, then carry the solution across.')

# ---------------------------------------------------- 1.  x -> 1/x
Qline(1, 'a', [t('Solve the inequality '),
               m(r'3\left|x^{2}-4\right| \leq |x+2|'), t('.')], marks=4)
Pline('b', [t('Hence solve '),
            m(r'\dfrac{3}{x^{2}}\left|1-4x^{2}\right| \leq \left|\dfrac{1+2x}{x}\right|'),
            t('.')], marks=2)
ws.ans([t('(a) '), m(r'\dfrac{5}{3} \leq x \leq \dfrac{7}{3}'), t(' or '), m('x=-2'),
        t('   (b) '), m(r'x \to \dfrac{1}{x}'), t(': '),
        m(r'\dfrac{3}{7} \leq x \leq \dfrac{3}{5}'), t(' or '), m(r'x=-\dfrac{1}{2}')])

# ---------------------------------------------------- 2.  x -> -x, reversed
Qline(2, 'a', [t('Solve the inequality '),
               m(r'\dfrac{x+5}{(x-3)^{2}} \geq \dfrac{x-7}{x(x-3)}'), t('.')], marks=4)
Pline('b', [t('Hence solve '),
            m(r'\dfrac{x-5}{(x+3)^{2}} \leq \dfrac{x+7}{x(x+3)}'), t('.')], marks=2)
ws.ans([t('(a) '), m('x<0'), t(' or '), m(r'x \geq \dfrac{7}{5}'), t(', '), m(r'x \neq 3'),
        t('   (b) '), m(r'x \to -x'), t(', then '), m(r'\times(-1)'), t(': '),
        m('x>0'), t(' or '), m(r'x \leq -\dfrac{7}{5}'), t(', '), m(r'x \neq -3')])

# ---------------------------------------------------- 3.  x -> -|x|
Qline(3, 'a', [t('Solve exactly '), m(r'\dfrac{x^{2}-x-1}{x+1} \leq 1'), t('.')], marks=4)
Pline('b', [t('Hence solve exactly '),
            m(r'\dfrac{x^{2}+|x|-1}{1-|x|} \leq 1'), t('.')], marks=4)
ws.ans([t('(a) '), m('x<-1'), t(' or '), m(r'1-\sqrt{3} \leq x \leq 1+\sqrt{3}'),
        t('   (b) '), m(r'x \to -|x|'), t(': '),
        m('x<-1'), t(' or '), m('x>1'), t(' or '),
        m(r'1-\sqrt{3} \leq x \leq \sqrt{3}-1')])

# ---------------------------------------------------- 4.  x -> sin x, hidden
# The target shows no sin x at all: cosec x and cos^2 x have to be rewritten
# as 1/sin x and 1 - sin^2 x, and the whole thing rearranged, before the
# shape of part (a) appears.
Qline(4, 'a', [t('Solve the inequality '),
               m(r'x - \dfrac{1}{x} \geq x^{2}-1'), t('.')], marks=4)
# "cosec" goes in as a plain text run: \operatorname and \csc both come out
# italic and \mathrm comes out letter-spaced, whereas a text run renders
# upright and tight in every renderer.
Pline('b', [t('Hence solve the inequality '),
            m(r'\sin x'), t(' \u2212 cosec '), m('x'), t(' + '),
            m(r'\cos^{2}x \geq 0'),
            t(' for '), m(r'0 \leq x \leq 2\pi'), t('.')], marks=3)
ws.ans([t('(a) '), m(r'-1 \leq x < 0'), t(' or '), m('x=1'),
        t('   (b) rewrite as '),
        m(r'\sin x - \dfrac{1}{\sin x} \geq \sin^{2}x - 1'),
        t(', then '), m(r'x \to \sin x'), t(': '),
        m(r'\pi < x < 2\pi'), t(' or '), m(r'x=\dfrac{\pi}{2}')])


# ---------------------------------------------------- 5.  x -> x/2, complement
# A scaling rather than a function, so there is no substituted expression to
# spot anywhere in the target — the 3 becoming 12 is the only clue (two halved
# factors give a divide-by-4). Then the inequality runs the other way, so the
# answer is the complement of the substituted range.
Qline(5, 'a', [t('Solve the inequality '),
               m(r'\dfrac{x}{x-2} < \dfrac{3}{(5x+2)(2-x)}'), t('.')], marks=3)
Pline('b', [t('Hence solve '),
            m(r'\dfrac{x}{x-4} > \dfrac{12}{(5x+4)(4-x)}'), t('.')], marks=2)
ws.ans([t('(a) '), m(r'-\dfrac{2}{5} < x < 2'),
        t('   (b) '), m(r'x \to \dfrac{x}{2}'), t(', then the complement: '),
        m(r'x < -\dfrac{4}{5}'), t(' or '), m('x>4')])

ws.save('inequalities_substitution.docx')
print('saved')
