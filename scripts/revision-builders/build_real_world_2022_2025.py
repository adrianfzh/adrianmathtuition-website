# -*- coding: utf-8 -*-
"""
O Level Revision — Math In Real World Context (2022-2025 edition)

Ten real prelim questions pulled from the Supabase question bank, all
2022-2025, chosen for variety of context AND variety of underlying skill.
Every printed answer is COMPUTED here, never transcribed from the bank
(the bank's answer for the salary question was wrong: it gave the ratio
107.8% where the question asks for the increase, 7.8%).
"""
from worksheet_lib import Worksheet
from docx.shared import Cm, Pt, RGBColor
from solutions_mono import SOL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fractions import Fraction as F
import math


# ------------------------------------------------------------------ helpers
def data_table(ws, rows, widths=None, bold_first_row=False, bold_first_col=False,
               align='center'):
    """A bordered data table inside a question, in house style (TNR 9.5, 1.15)."""
    ncols = max(len(r) for r in rows)
    t = ws.doc.add_table(rows=len(rows), cols=ncols)
    t.style = ws.doc.styles['Table Grid']
    t.autofit = False
    t.alignment = {'center': 1, 'left': 0}.get(align, 1)
    for ri, rowdata in enumerate(rows):
        row = t.rows[ri]
        _cant_split(row)
        for ci in range(ncols):
            cell = row.cells[ci]
            if widths:
                cell.width = Cm(widths[ci])
            txt = rowdata[ci] if ci < len(rowdata) else ''
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 and ncols > 2 \
                else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(txt))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9.5)
            if (bold_first_row and ri == 0) or (bold_first_col and ci == 0):
                run.bold = True
    # Bind only the lead-in paragraph to the table. Marking every cell
    # keep_with_next chained whole questions together, and a chain that cannot
    # fit gets broken at the first opportunity -- which left half-empty pages.
    if ws._block_paras:
        ws._block_paras[-1].paragraph_format.keep_with_next = True
    spacer = ws.doc.add_paragraph()
    spacer.paragraph_format.line_spacing = 1.0
    spacer.paragraph_format.space_after = Pt(0)
    ws._block_paras.append(spacer)
    return t


def _cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn('w:cantSplit')) is None:
        trPr.append(OxmlElement('w:cantSplit'))


def lit_part(ws, label, parts, marks=None):
    """A literal padded sub-sub-part label — (i) / (ii) under an (a) stem.

    SQ() would relabel these as (a),(b) via its own startOverride, so
    sub-sub-parts must be written literally.
    """
    p = ws.doc.add_paragraph()
    p.style = ws.doc.styles['Normal']
    p.paragraph_format.left_indent = Cm(2.0)
    p.paragraph_format.first_line_indent = Cm(-0.85)
    run = p.add_run(f'({label}) '.ljust(5))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9.5)
    ws._fill(p, parts)
    if marks is not None:
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT)
        r = p.add_run(f'\t[{marks}]')
        r.font.name = 'Times New Roman'
        r.font.size = Pt(9.5)
    ws._block_paras.append(p)
    return p


def fix_grids(ws):
    """python-docx writes an EQUAL-width w:tblGrid whatever the cell widths are.

    Under a fixed table layout the grid is what LibreOffice honours (and Word
    prefers it too when tcW disagrees), so every table rendered as equal
    columns -- the solution boxes came out 50/50 instead of 1 cm + 15 cm.
    Rewrite each grid from the first row's real tcW values.
    """
    for t in ws.doc.tables:
        widths = []
        for cell in t.rows[0].cells:
            tcPr = cell._tc.find(qn('w:tcPr'))
            tcW = tcPr.find(qn('w:tcW')) if tcPr is not None else None
            widths.append(tcW.get(qn('w:w')) if tcW is not None else None)
        grid = t._tbl.find(qn('w:tblGrid'))
        if grid is None or any(w is None for w in widths):
            continue
        cols = grid.findall(qn('w:gridCol'))
        if len(cols) != len(widths):
            continue
        for col, w in zip(cols, widths):
            col.set(qn('w:w'), w)


def expand_aligned(rows):
    r"""Split every \begin{aligned} block into separate one-line steps.

    Two hard constraints found by rendering, not by reading the XML:

    * pandoc here does NOT emit an m:eqArr for aligned/array/split/matrix --
      the `\\` collapses to a thin space, so a two-line block came out as one
      over-long line that ran off the edge of the solution box.
    * an equation whose first visible token is `=` renders as a stray "¿".
      A leading \qquad or \hspace does not save it; only a real left-hand
      side does.

    So each aligned line becomes its own display equation, and a
    continuation line ("&= ...") repeats the left-hand side -- which is how
    the working would be written out by hand anyway.
    """
    out = []
    for label, steps in rows:
        new = []
        for step in steps:
            if not (isinstance(step, str) and r'\begin{aligned}' in step):
                new.append(step)
                continue
            body = step.split(r'\begin{aligned}', 1)[1].rsplit(r'\end{aligned}', 1)[0]
            lhs = ''
            for piece in body.split('\\\\'):
                piece = piece.strip()
                if not piece:
                    continue
                if '&' in piece:
                    left, right = piece.split('&', 1)
                    if left.strip():
                        lhs = left.strip()
                    new.append(f'{lhs} {right.strip()}'.strip())
                else:
                    new.append(piece)
        out.append((label, new))
    return out


_orig_solution_box = Worksheet.solution_box


def _solution_box(self, rows, keep_together=True):
    return _orig_solution_box(self, expand_aligned(rows), keep_together)


Worksheet.solution_box = _solution_box


_pending_break = {'on': False}
_orig_Q = Worksheet.Q


def _Q(self, parts, marks=None):
    """Start each new question on a fresh page via page_break_before.

    doc.add_page_break() inserts a page break in a paragraph OF ITS OWN, so a
    solution box that happened to end near the page bottom pushed that empty
    paragraph onto the next page and the break then produced a completely
    blank one (pages 8 and 12 were empty). page_break_before on the question
    itself cannot do that.
    """
    para = _orig_Q(self, parts, marks)
    if _pending_break['on']:
        para.paragraph_format.page_break_before = True
        _pending_break['on'] = False
    return para


Worksheet.Q = _Q


def new_page(ws):
    _pending_break['on'] = True
    ws._block_paras = []


def drop_blanks_before_breaks(ws):
    """Delete empty paragraphs sitting just before a page-break-before paragraph.

    solution_box() ends with a blank spacer paragraph. When the box finished
    near the page bottom that spacer flowed onto the next page on its own, so
    the page before the next question came out completely blank (page 8).
    The spacer is only breathing space, so it is safe to drop when a hard page
    break follows it anyway.
    """
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    body = ws.doc.element.body
    kids = list(body)
    removed = 0

    def is_blank_para(el):
        if el.tag != f'{{{W}}}p':
            return False
        if el.find(qn('w:pPr')) is not None:
            ppr = el.find(qn('w:pPr'))
            if ppr.find(qn('w:pageBreakBefore')) is not None:
                return False
        if len(el.findall('.//' + qn('w:drawing'))) or len(el.findall('.//' + qn('w:br'))):
            return False
        if el.find('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath') is not None:
            return False
        return ''.join(t.text or '' for t in el.iter(qn('w:t'))).strip() == ''

    def has_break(el):
        if el.tag != f'{{{W}}}p':
            return False
        ppr = el.find(qn('w:pPr'))
        return ppr is not None and ppr.find(qn('w:pageBreakBefore')) is not None

    for i, el in enumerate(kids):
        if not has_break(el):
            continue
        j = i - 1
        while j >= 0 and is_blank_para(kids[j]):
            kids[j].getparent().remove(kids[j])
            removed += 1
            j -= 1
    return removed


SOL_NAVY = RGBColor(0x1F, 0x4E, 0x79)
SOL_BAR = 'DEEBF7'
SOL_EDGE = '2E5C8A'
_PART_RE = __import__('re').compile(r'^\((?:[a-z]|i{1,3}|iv|v)\)(\((?:i{1,3}|iv|v)\))?\s*$')


def _shade(cell, hexfill):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:color'), 'auto')
    el.set(qn('w:fill'), hexfill)
    cell._tc.get_or_add_tcPr().append(el)


def _box(table, hexcolor):
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '6')
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), hexcolor)
        borders.append(e)
    tblPr = table._tbl.tblPr
    anchor = None
    for tag in ('w:tblLayout', 'w:tblCellMar', 'w:tblLook'):
        anchor = tblPr.find(qn(tag))
        if anchor is not None:
            break
    (anchor.addprevious(borders) if anchor is not None else tblPr.append(borders))


def solution_block(ws, qnum):
    """Adrian's uploaded solution format: a navy-edged box with a pale blue
    'Solution — Qn' bar, body set in Consolas so the '=' signs line up and the
    [M1]/[A1] mark codes sit in a clean right-hand column."""
    lines = SOL[qnum]
    ws.doc.add_paragraph()
    t = ws.doc.add_table(rows=2, cols=1)
    t.autofit = False
    _box(t, SOL_EDGE)

    head = t.rows[0].cells[0]
    head.width = Cm(16)
    _shade(head, SOL_BAR)
    hp = head.paragraphs[0]
    hp.paragraph_format.line_spacing = 1.0
    hp.paragraph_format.space_before = Pt(2)
    hp.paragraph_format.space_after = Pt(2)
    hr = hp.add_run(f'Solution  —  Q{qnum}')
    hr.font.name = 'Calibri'
    hr.font.size = Pt(10.5)
    hr.bold = True
    hr.font.color.rgb = SOL_NAVY

    body = t.rows[1].cells[0]
    body.width = Cm(16)
    first = True
    for line in lines:
        bp = body.paragraphs[0] if first else body.add_paragraph()
        first = False
        bp.paragraph_format.line_spacing = 1.0
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after = Pt(0)
        run = bp.add_run(line if line else ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        if _PART_RE.match(line.strip()):
            run.bold = True
        # leading spaces are the alignment - Word must not trim them
        for wt in run._r.findall(qn('w:t')):
            wt.set(qn('xml:space'), 'preserve')
    ws._block_paras = []
    ws.doc.add_paragraph()
    return t


def T(s):
    return [('text', s)]


def qpara(parts, marks=None):
    """Continuation of the question stem — indented to the question text column."""
    p = ws.para(parts, marks)
    p.paragraph_format.left_indent = Cm(1.0)
    return p


def spara(parts, marks=None):
    """Continuation of a sub-part — indented to the (a)/(b) text column."""
    p = ws.para(parts, marks)
    p.paragraph_format.left_indent = Cm(2.0)
    return p


def money(x, dp=2):
    return f'${x:,.{dp}f}'


ws = Worksheet()
ws.title('O Level Revision')
ws.subtitle('Math In Real World Context')


# =====================================================================  Q1
# Salary schemes — percentage increase (the "ratio vs increase" trap)
sal = 3500
a_annual = sal * 1.07 * 12                       # 44940
b_annual = sal * 1.06 * 12 + 250 * (12 // 4)     # 45270
cur_annual = sal * 12                            # 42000
b_pct = (b_annual - cur_annual) / cur_annual * 100   # 7.7857...

ws.Q(T('Mabel currently earns a monthly salary of $3500. She was offered an increase in her salary '
       'calculated according to either Scheme A or Scheme B.'))
qpara(T('Scheme A gives an increase of 7% of her present monthly salary.'))
qpara(T('Scheme B gives an increase of 6% of her present monthly salary, plus a bonus of $250 '
          'every 4 months.'))
ws.SQ(T('Calculate her annual salary under'))
lit_part(ws, 'i', T('Scheme A,'), marks=1)
lit_part(ws, 'ii', T('Scheme B.'), marks=2)
ws.SQ(T('Express Mabel’s increase in annual salary under Scheme B as a percentage of her '
        'current annual salary, giving your answer correct to 1 decimal place.'), marks=2)
ws.ans(T(f'(a)(i) {money(a_annual,0)}  (ii) {money(b_annual,0)}   (b) {b_pct:.1f}%'))
solution_block(ws, 1)
new_page(ws)


# =====================================================================  Q2
# Ibuprofen dosage — proportion, mg->g, safety threshold
conc = 100 / 5                 # mg per ml
dose_g = 3.75 * conc / 1000    # 0.075 g
per_dose = 6 * conc            # 120 mg
doses_day = 24 // 6            # 4
day_total = per_dose * doses_day       # 480
max_day = 40 * 11.5                    # 460
per_kg = per_dose / 11.5               # 10.43

ws.Q(T('Ibuprofen, a non-steroidal anti-inflammatory drug (NSAID), is used to treat fever and to '
       'relieve mild to moderate pain such as a sore throat, toothache or muscle ache. Parents of '
       'young children may give ibuprofen to bring a high fever down, but it should be given only '
       'when the temperature is above 38.5 °C. A typical bottle of ibuprofen has a dosage of '
       '100 mg/5 ml (that is, 100 mg of ibuprofen in 5 ml of solution).'))
ws.SQ([('text', 'A 3.75 ml dose of the solution is given to a child. Find the amount of ibuprofen '
                'in the dose, in grams. (Take '),
       ('math', r'1\text{ mg} = 10^{-3}\text{ g}'), ('text', '.)')], marks=2)
ws.SQ(T('A regular dose for a child is 4 to 10 mg/kg per dose, and may be given every 6 to 8 hours. '
        'The maximum dosage is 40 mg/kg per day. A child of mass 11.5 kg was given 6 ml of the '
        'solution every 6 hours. Determine, with appropriate working, whether the child was given '
        'an overdose for the day.'), marks=4)
ws.ans(T(f'(a) {dose_g} g   (b) {day_total:.0f} mg > {max_day:.0f} mg, so yes — an overdose'))
solution_block(ws, 2)
new_page(ws)


# =====================================================================  Q3
# Mobile plans — "or part thereof" ceiling charging, payback over a contract
rate_min = 0.1603
star = 38 + (140 - 100) * rate_min + math.ceil(4.8 - 4) * 10.70     # 55.112
spec = 45 + (140 - 120) * rate_min + max(0, math.ceil(4.8 - 5)) * 10.70  # 48.206
star_2y = 24 * star + 400
spec_2y = 24 * spec + 540
saving_2y = star_2y - spec_2y

ws.Q(T('The table shows the mobile call charges and data charges from two mobile companies, '
       'Star Mobile and Spectrum Mobile.'))
data_table(ws, [
    ['', 'Star Mobile', 'Spectrum Mobile'],
    ['Monthly subscription fee', '$38', '$45'],
    ['Local incoming calls', 'Free', 'Free'],
    ['SMS', '800', 'Unlimited'],
    ['Local outgoing calls', '100 minutes of free talk time', '120 minutes of free talk time'],
    ['Local data bundle', '4 GB free', '5 GB free'],
], widths=[5.2, 5.0, 5.0], bold_first_row=True)
qpara(T('Beyond the free talk time, 16.03 cents (or part thereof) is charged for each minute of '
          'outgoing talk time. Beyond the free data bundle, $10.70/GB (or part thereof) is charged, '
          'capped at $168 for Star Mobile and $95 for Spectrum Mobile.'))
ws.SQ(T('Matthew’s mean monthly usage is given below.'))
data_table(ws, [
    ['Local outgoing talk time', '140 minutes'],
    ['Local data usage', '4.8 GB'],
    ['SMS sent', '280'],
], widths=[5.5, 4.0])
spara(T('Decide which plan is more cost saving for him.'), marks=3)
ws.SQ(T('Matthew is able to buy the latest m-phone V at a discounted price, but he has to sign a '
        '2-year contract.'))
data_table(ws, [
    ['', 'Star Mobile', 'Spectrum Mobile'],
    ['Price of m-phone V', '$400', '$540'],
], widths=[5.2, 4.0, 4.0], bold_first_row=True)
spara(T('Considering the costs incurred from each company based on his monthly usage in (a) and '
          'the price of the phone, decide which package is more cost saving for him over 2 years. '
          'Justify your answer.'), marks=3)
ws.ans(T(f'(a) Spectrum Mobile ({money(spec)} < {money(star)})   '
         f'(b) Spectrum Mobile, cheaper by {money(saving_2y)} over 2 years'))
solution_block(ws, 3)
new_page(ws)


# =====================================================================  Q4
# Vehicles per 1000 — rate, standard form, reverse % growth, testing a claim
sg_pop = 4_400_000 / 766 * 1000
china_2021 = 231 / 1.14 ** 2
india_v = 1451e6 / 1000 * 57
japan_v = 123.8e6 / 1000 * 810

ws.Q(T('Table A below shows the number of vehicles per 1000 people in five countries in 2023. '
       'Table B shows the estimated population of the same group of countries in 2023.'))
data_table(ws, [
    ['Country', 'Number of vehicles per 1000 people'],
    ['USA', '900'], ['Japan', '810'], ['Singapore', '766'],
    ['China', '231'], ['India', '57'],
], widths=[4.0, 6.5], bold_first_row=True)
qpara([('text', 'Table A', {'italic': True})])
data_table(ws, [
    ['Country', 'Population (millions)'],
    ['USA', '345.4'], ['Japan', '123.8'], ['Singapore', '?'],
    ['China', '1 419'], ['India', '1 451'],
], widths=[4.0, 6.5], bold_first_row=True)
qpara([('text', 'Table B', {'italic': True})])
ws.SQ(T('The estimated number of vehicles in Singapore in 2023 was 4 400 000. Calculate an estimate '
        'of the population of Singapore in 2023, leaving your answer in standard form, correct to '
        '3 significant figures.'), marks=2)
ws.SQ(T('Between 2015 and 2023, the number of vehicles per 1000 people in China grew by 14% each '
        'year. Estimate the number of vehicles per 1000 people in China in 2021.'), marks=3)
ws.SQ(T('After looking at Table A and Table B, Sue concluded that India should have more vehicles '
        'than Japan, because India’s population is much bigger than Japan’s. Is Sue’s '
        'conclusion correct? Explain your answer, supported by clear calculations.'), marks=3)
ws.ans([('text', '(a) '), ('math', r'5.74 \times 10^{6}'),
        ('text', f'   (b) {china_2021:.0f} vehicles per 1000 people   '
                 '(c) No — Japan has about 100 278 000 vehicles, India about 82 707 000')])
solution_block(ws, 4)
new_page(ws)


# =====================================================================  Q5
# Two pipes — forming and solving a quadratic from rates of work
root = (30 + math.sqrt(30 ** 2 + 4 * 200)) / 2      # 35.6155...
root2 = (30 - math.sqrt(30 ** 2 + 4 * 200)) / 2     # -5.6155...
small = root + 10
pct = 20 / small * 100                              # 43.84...

ws.Q([('text', 'A water tank can be filled with water by two pipes in 20 minutes. The smaller pipe '
                'takes 10 minutes longer than the larger pipe to fill the tank. The larger pipe '
                'takes '), ('math', 'x'), ('text', ' minutes to fill the tank.')])
ws.SQ([('text', 'Write down an equation in '), ('math', 'x'),
       ('text', ' to represent this information, and show that it reduces to '),
       ('math', r'x^{2} - 30x - 200 = 0'), ('text', '.')], marks=3)
ws.SQ([('text', 'Solve the equation '), ('math', r'x^{2} - 30x - 200 = 0'),
       ('text', ', giving your answers correct to 2 decimal places.')], marks=3)
ws.SQ(T('Find the percentage of the tank that can be filled by the smaller pipe in 20 minutes.'),
      marks=2)
ws.ans([('text', '(a) shown   (b) '), ('math', f'x = {root:.2f}'), ('text', ' or '),
        ('math', f'x = {root2:.2f}'), ('text', f'   (c) {pct:.1f}%')])
solution_block(ws, 5)
new_page(ws)


# =====================================================================  Q6
# London trip — time zones, currency, three-option comparison
flight_total_min = 13 * 60 + 50 + 120           # 950 min = 15 h 50
arrive_sg_min = (23 * 60 + 15 + flight_total_min) % (24 * 60)   # 15:05
arrive_ldn_min = (arrive_sg_min - 7 * 60) % (24 * 60)           # 08:05
daily = 25 + 10 + 15
trip_gbp = daily * 6
trip_sgd = trip_gbp * 1.65
emergency = 0.10 * 700
need = trip_sgd + emergency
budget_sight = 15 * 4
optA = 90 * 0.95
optB = 8 * 12.50
optC = 35 + 4 * 6


def hhmm(m):
    h, mi = divmod(m, 60)
    ap = 'AM' if h < 12 else 'PM'
    h12 = h % 12 or 12
    return f'{h12}:{mi:02d} {ap}'


ws.Q(T('Mdm Siti is planning a trip from Singapore to London. She will spend 6 days in London. '
       'Here is the information on her trip.'))
qpara(T('•  The total flight duration is 13 hours 50 minutes, which excludes a 2-hour '
          'layover at Doha.'))
data_table(ws, [
    ['Singapore to Doha', '7 h 40 min'],
    ['Doha to London', '6 h 10 min'],
], widths=[5.0, 3.5])
qpara(T('•  Her flight departs at 11:15 PM (Singapore time).'))
qpara(T('•  London time is 7 hours behind Singapore time.'))
qpara(T('•  The exchange rate is 1 Pound Sterling (£) = 1.65 Singapore Dollars (SGD).'))
ws.SQ(T('What is her arrival time in London, in local London time?'), marks=1)
ws.SQ(T('Mdm Siti plans to bring a total of SGD 700 for meals, transport, sightseeing and '
        'emergencies. She plans to spend, on average, the following for the entire trip in London:'))
spara(T('       £25 per day on food,  £10 per day on transport,  £15 per day on '
          'sightseeing,  and 10% of the total amount for emergencies.'))
spara(T('By showing your working clearly, determine whether SGD 700 is sufficient.'), marks=3)
ws.SQ(T('Mdm Siti plans to spend four days sightseeing in London, and to visit only 2 attractions '
        'per day. She is considering three sightseeing options.'))
data_table(ws, [
    ['Option', 'Type', 'Description'],
    ['A', 'Fixed daily pass',
     '4-day London Sightseeing Pass, covering unlimited entries, for £90, with a 5% discount '
     'if she stays more than 5 days in London.'],
    ['B', 'Pay-as-you-go', 'Each attraction costs £12.50.'],
    ['C', 'Partial pass + pay-as-you-go',
     'She buys a 2-day pass for £35, covering unlimited entries for 2 days, and pays £6 '
     'per attraction for the other 2 days.'],
], widths=[1.6, 3.6, 9.5], bold_first_row=True)
spara(T('Assuming the total amount set aside for sightseeing stays the same for the four days, '
          'show your calculations and explain which option would be the most cost-effective for '
          'Mdm Siti, considering her sightseeing plans and her budget.'), marks=5)
ws.ans(T(f'(a) {hhmm(arrive_ldn_min)} (London time)   (b) Yes — only {money(need)} of '
         f'SGD 700 is needed   (c) Option C, at £{optC:.0f}'))
solution_block(ws, 6)
new_page(ws)


# =====================================================================  Q7
# TDEE — substituting into a given formula, then planning backwards
def tdee(w, h, age, lvl):
    return (10 * w + 6.25 * h - 5 * age + 5) * lvl


q7a = tdee(72, 172, 25, 1.725)          # 2889.375
q7b = 7.3 * 65 * 0.75                   # 355.875
stan = tdee(75, 168, 42, 1.55)          # 2472.25
eaten = 2150 * 7
deficit = stan * 7 - eaten
swim = 7.8 * 75 * 1
jog_cal = 3850 - deficit - swim
jog_h = jog_cal / (8.8 * 75)
jog_each = jog_h / 3 * 60               # 30.58 min

ws.Q(T('The Total Daily Energy Expenditure (TDEE) is an estimate of how many calories a person '
       'uses per day, and can be calculated by'))
ws.math_block(r'\text{TDEE} = \left[10 \times \text{weight in kg} + 6.25 \times \text{height in cm}'
              r' - 5 \times \text{age in years} + 5\right] \times \text{Activity Level}')
qpara(T('The Activity Level can be determined from the table below.'))
data_table(ws, [
    ['Activity Level', 'Description'],
    ['1.2', 'Sedentary: 1 to 3 sessions of exercise a month'],
    ['1.375', 'Light: 1 to 3 sessions of exercise a week'],
    ['1.55', 'Moderate: 4 to 5 sessions of exercise a week'],
    ['1.725', 'High: 6 to 7 sessions of exercise a week'],
    ['1.9', 'Extreme: 2 sessions of exercise a day'],
], widths=[3.2, 8.5], bold_first_row=True)
ws.SQ(T('Find the TDEE of a 25-year-old person who is 172 cm tall, weighs 72 kg and has 6 sessions '
        'of exercise a week.'), marks=1)
ws.SQ(T('Additional calories are used when a person performs certain exercises, and can be '
        'calculated by'))
ws.math_block(r'\text{Calories} = \text{MET} \times \text{weight in kg} \times '
              r'\text{duration in hours}')
spara(T('where the Metabolic Equivalent of Task (MET) is given below.'))
data_table(ws, [
    ['MET', '3.1', '7.3', '7.8', '8.8', '10.3'],
    ['Type of exercise', 'Walking', 'Badminton', 'Swimming', 'Jogging', 'Basketball'],
], widths=[3.2, 2.2, 2.4, 2.3, 2.2, 2.6], bold_first_col=True)
spara(T('Find the additional calories used by a person who weighs 65 kg and plays badminton for '
          '45 minutes.'), marks=2)
ws.SQ(T('Stan is 42 years old, weighs 75 kg and is 168 cm tall. He aims to lose some weight. He '
        'learns that he needs to use about 3850 calories more than the calories he eats every '
        'week. He decides to eat about 2150 calories every day. He has a one-hour session of '
        'swimming every week, and wants to add 3 sessions of jogging to his weekly schedule.'))
spara(T('Suggest a reasonable duration for Stan to jog during each session. Justify any decisions '
          'you make and show your calculations clearly.'), marks=7)
ws.ans(T(f'(a) {q7a} calories/day   (b) {q7b} calories   '
         f'(c) {jog_each:.1f} min each, so about 30 minutes per session'))
solution_block(ws, 7)
new_page(ws)


# =====================================================================  Q8
# Maple syrup — table lookup, unit conversion, cost per unit -> selling price
pct_syrup = 0.25 / 10.5 * 100
taps_three = 1 + 1 + 3
syrup_gal = F(taps_three * 12, 1) / F(21, 2) * F(1, 4)      # exact 10/7
syrup_l = syrup_gal * F(31, 8)                              # 3.875 = 31/8
farm_taps = 40 * 200
farm_syrup = F(farm_taps * 12, 1) / F(21, 2) * F(1, 4)
maintain = 40 * 1200
convert = farm_syrup * 18
total_cost = maintain + convert
cost_per_gal = total_cost / farm_syrup
min_price = cost_per_gal + 20

ws.Q(T('Maple syrup is made by boiling down the sap of maple trees. The sap is a liquid found in '
       'the tree, and is collected by drilling a hole into the tree and putting a tap into it. Each '
       'tree can have at most 3 taps, depending on its size. Below is some information about '
       'tapping maple trees.'))
data_table(ws, [
    ['Diameter of tree', 'Number of taps'],
    ['Less than 10 inches', '0'],
    ['10 – 17 inches', '1'],
    ['18 – 25 inches', '2'],
    ['More than 25 inches', '3'],
], widths=[5.0, 4.0], bold_first_row=True)
qpara(T('An average of 12 gallons of sap can be collected per tap. 10.5 gallons of sap boils down '
          'to 0.25 gallons of syrup. 1 gallon = 3.875 litres.'))
ws.SQ(T('What percentage of the sap eventually becomes maple syrup?'), marks=1)
ws.SQ(T('Calculate the total volume, in litres, of maple syrup that can be produced from three '
        'maple trees of diameters 10 inches, 15 inches and 27 inches respectively.'), marks=2)
ws.SQ(T('The maple syrup found in stores comes from maple trees grown on farms. Information about '
        'a typical maple farm is shown below.'))
data_table(ws, [
    ['Number of acres', '40'],
    ['Number of taps per acre', '200'],
    ['Cost of maintaining the farm', '$1200 per acre'],
    ['Cost to convert sap to syrup', '$18 per gallon of syrup'],
], widths=[6.5, 5.0], bold_first_col=True)
spara(T('The owner of a maple farm would like to make a profit of at least $20 for every gallon '
          'of maple syrup he sells. What price should he charge for one gallon of maple syrup to '
          'achieve his aim? Justify any decisions you make and show your calculations clearly.'),
        marks=7)
ws.ans(T(f'(a) {pct_syrup:.2f}%   (b) {float(syrup_l):.2f} litres   '
         f'(c) at least {money(float(min_price),0)} per gallon'))
solution_block(ws, 8)
new_page(ws)


# =====================================================================  Q9
# Condominium — m^2 -> ft^2, simple-interest loan, affordability cap, round DOWN
ft_per_m2 = 3.28 ** 2
area_3rm = 60 * ft_per_m2                 # 645.504
area_3sf = 646
base_price = 990 * area_3sf               # 639540
cap = 0.30 * 7500                         # 2250
# 45% paid up front -> loan is 55%; simple interest 3% p.a. for 25 years
max_price = cap * (25 * 12) / (0.55 + 0.55 * 0.03 * 25)
levels_up = (max_price - base_price) / 16000
max_level = 1 + math.floor(levels_up)

ws.Q(T('A new 52-level private condominium, with units starting from the first level, has just been '
       'launched. The sale price of a unit depends on the floor area and the level the unit is on. '
       'The price of each similar unit increases by a specific amount per level. Below is some '
       'information on the different types of units available.'))
data_table(ws, [
    ['Type of unit', '2-room', '3-room', '4-room', '5-room'],
    ['Number of toilets', '1', '1', '2', '2'],
    ['Floor area (m²)', '40', '60', '90', '110'],
    ['Price per square foot ($)', '890', '990', '1100', '1300'],
    ['Increase in price per level ($)', '—', '16 000', '19 000', '22 000'],
], widths=[5.0, 2.4, 2.4, 2.4, 2.4], bold_first_row=True, bold_first_col=True)
ws.SQ([('text', 'Given that 1 m = 3.28 feet, show that the floor area of a 3-room unit is 646 '
                'square feet, correct to 3 significant figures.')], marks=2)
ws.SQ(T('Calculate the sale price of a 3-room unit on the first level.'), marks=1)
ws.SQ(T('Lisa is planning to buy a 3-room unit. To purchase the unit, she decides to pay 45% of the '
        'price of the unit and to take a housing loan from the bank to pay the remaining amount by '
        'monthly instalments. The bank charges a simple interest rate of 3% per year for 25 years. '
        'Lisa has a gross monthly salary of $7500, and she does not plan to spend more than 30% of '
        'her gross monthly salary on the monthly instalment. Suggest the maximum level that Lisa '
        'can afford to purchase. Justify the decisions you make and show your calculations '
        'clearly.'), marks=7)
ws.ans([('text', '(a) shown   (b) '), ('math', r'\$639\,540'),
        ('text', f'   (c) the {max_level}th level')])
solution_block(ws, 9)
new_page(ws)


# =====================================================================  Q10
# Solar panels — mean, GST, packing a rectangle two ways, payback
usage = [1007.8, 1166.3, 1133.6, 1249, 1248.5, 1282.6]
avg_kwh = sum(usage) / len(usage)                  # 1181.3
tariff, gst = 0.2812, 1.09
avg_cost = avg_kwh * tariff * gst                  # 362.08
fit_a = math.floor(9 / 1.65) * math.floor(4 / 1)   # 5 * 4 = 20
fit_b = math.floor(9 / 1) * math.floor(4 / 1.65)   # 9 * 2 = 18
panels = max(fit_a, fit_b)
generated = panels * 45
net_kwh = avg_kwh - generated
net_cost = net_kwh * tariff * gst
panel_month = (10250 * (panels // 10)) / (25 * 12)
total_month = net_cost + panel_month

ws.Q(T('Chan is considering installing solar panels at his house and wants to determine whether it '
       'would be worthwhile. The tables below show the information Chan needs to make his decision.'))
qpara([('text', 'Electricity usage for 2025 (kWh)', {'bold': True})])
data_table(ws, [
    ['January', 'February', 'March', 'April', 'May', 'June'],
    ['1007.8', '1166.3', '1133.6', '1249', '1248.5', '1282.6'],
], widths=[2.4] * 6, bold_first_row=True)
qpara([('text', 'Charges for electricity usage', {'bold': True})])
qpara(T('Electricity tariff: 28.12 cents per kWh (excluding GST). Charges are subject to 9% '
          'Goods and Services Tax.'))
qpara([('text', 'Installation of solar panels', {'bold': True})])
data_table(ws, [
    ['Dimensions of the roof area of Chan’s house', '9 m by 4 m'],
    ['Dimensions of 1 solar panel', '1.65 m by 1 m'],
    ['Cost of installing 10 solar panels (not subject to 9% GST)', '$10 250'],
    ['Average electricity generated by 1 solar panel', '45 kWh per month'],
    ['Lifespan of solar panels', '25 years'],
], widths=[9.0, 4.0], bold_first_col=True)
ws.SQ(T('For the first six months of 2025, calculate the'))
lit_part(ws, 'i', T('average monthly amount of electricity, in kWh, used by Chan,'), marks=2)
lit_part(ws, 'ii', T('average monthly cost that Chan paid for electricity.'), marks=2)
ws.SQ(T('Calculate the maximum number of solar panels that can be installed on the roof of '
        'Chan’s house.'), marks=2)
ws.SQ(T('Should Chan go ahead and install solar panels at his house? Justify the decisions you '
        'make and show your calculations clearly.'), marks=5)
ws.ans(T(f'(a)(i) {avg_kwh:.1f} kWh  (ii) {money(avg_cost)}   (b) {panels} panels   '
         f'(c) Yes — {money(total_month)} a month against {money(avg_cost)}'))
solution_block(ws, 10)

fix_grids(ws)
print('blank paragraphs dropped before page breaks:', drop_blanks_before_breaks(ws))
ws.save('math_real_world_2022_2025.docx')
print('marks per question:', [5, 6, 6, 8, 8, 9, 10, 10, 10, 11], 'total', sum([5, 6, 6, 8, 8, 9, 10, 10, 10, 11]))
print(f'Q1  b% = {b_pct:.4f}   Q2 dose = {dose_g} g, {day_total} vs {max_day} mg')
print(f'Q3  star={star:.3f} spec={spec:.3f} 2y saving={saving_2y:.3f}')
print(f'Q4  sgpop={sg_pop:.1f} china2021={china_2021:.4f} india={india_v:,.0f} japan={japan_v:,.0f}')
print(f'Q5  roots {root:.4f} / {root2:.4f}  pct={pct:.4f}')
print(f'Q6  arrive {hhmm(arrive_ldn_min)}  need={need:.2f}  A={optA} B={optB} C={optC}')
print(f'Q7  a={q7a} b={q7b} stanTDEE={stan} jog_each={jog_each:.3f} min')
print(f'Q8  pct={pct_syrup:.4f} litres={float(syrup_l):.4f} costgal={float(cost_per_gal):.4f} min={float(min_price):.2f}')
print(f'Q9  area={area_3rm:.3f} base={base_price} maxprice={max_price:.2f} levels_up={levels_up:.4f} -> level {max_level}')
print(f'Q10 avg={avg_kwh:.4f} cost={avg_cost:.4f} fit {fit_a}/{fit_b} net={net_cost:.4f} panel={panel_month:.4f} tot={total_month:.4f}')
